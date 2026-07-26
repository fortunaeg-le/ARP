import re
import sys
import unicodedata
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.opc.exceptions import PackageNotFoundError
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.hyperlink import Hyperlink
from docx.oxml.ns import qn

from models import SourceDocument, TextSegment


# --- Нормализация регистра для detection_text (см. docs/SHIFRATOR_SPEC_AI.md, блок 1) ---
#
# Natasha опирается на регистр как признак имени собственного. Если текст .docx
# визуально заглавный через ФОРМАТИРОВАНИЕ (run.font.all_caps/small_caps), в XML
# физически лежат строчные буквы, python-docx их и возвращает — и ФИО в преамбуле
# договора остаётся в открытом виде. Рядом с настоящим segment.text заводим
# detection_text: копию ТОЙ ЖЕ ДЛИНЫ с нормализованным регистром. Детекторы ищут в
# detection_text, а original_text вырезают из настоящего text по тем же оффсетам —
# равная длина делает оффсеты валидными в обоих. Инвариант равной длины —
# фундамент схемы; при его нарушении detection_text НЕ создаётся вообще.

# Порог длины сегмента для эвристики Изменения 2 (настоящий строчный/заглавный ввод
# без форматирования). Узко — чтобы эвристика не превратилась во второй полный проход
# NER по всему документу.
_MAX_HEURISTIC_LEN = 200

# Признак пункта перечисления: "(1)" или "1." после ведущих пробелов. До этапа 2b
# это было УСЛОВИЕМ срабатывания эвристики, и в том и состояла «известная
# непокрытая дыра»: ФИО строчными в обычном абзаце не находилось. Этап 2b снял
# требование для обычных документов, но СОХРАНИЛ его для combo (см.
# `_apply_lowercase_heuristic`): там 2b обязан не добавлять ничего нового, но и не
# отнимать того, что работало до него.
_LIST_ITEM_RE = re.compile(r"^\s*(\(\d+\)|\d+\.)")


def _normalize_word(word: str) -> str:
    """Нормализует регистр ОДНОГО слова (максимального alpha-рана) посимвольно.

    НЕ используем str.title(): он (1) не гарантирует сохранение длины на лигатурах
    ('ﬁ' -> 'Fi') — а равная длина есть фундамент всей схемы; (2) ломает смысл —
    'ип пирогова' -> 'Ип Пирогова', разрушая юр-маркер 'ИП', на котором строится
    синтаксический слой составных сущностей ('ИП + ФИО') из следующей задачи.

    Вместо этого: слово длиной 2-3 буквы целиком поднимаем в верхний регистр —
    так 'ип'->'ИП', 'ооо'->'ООО', 'ао'->'АО', 'зао'->'ЗАО', 'пао'->'ПАО',
    'кфх'->'КФХ' переживают нормализацию all_caps-рана как юр-формы, а не как
    'Ип'/'Ооо'. Это часть посимвольной логики (правило по длине), а не список
    исключений: любое короткое (2-3) полностью буквенное слово в визуально
    заглавном контексте логично оставить заглавным. Остальные слова: первая буква
    в верхний, прочие в нижний — но ТОЛЬКО через .upper()/.lower() над отдельными
    символами (не над строкой целиком). Инициалы ('а' в 'а.с.') — слово длины 1,
    первая буква -> 'А'.

    Считаем и нумеруем БУКВЫ, а не символы: слово может нести внутри невидимые
    (SHY/ZWSP), и они не должны ни попадать в счёт длины 2-3, ни занимать место
    «первой буквы». Сами невидимые копируются как есть (длина сохраняется)."""
    n_letters = sum(1 for ch in word if ch.isalpha())
    if n_letters in (2, 3):
        return "".join(ch.upper() for ch in word)
    out, seen = [], 0
    for ch in word:
        if not ch.isalpha():
            out.append(ch)
            continue
        out.append(ch.upper() if seen == 0 else ch.lower())
        seen += 1
    return "".join(out)


def _normalize_case(text: str) -> str:
    """Нормализует регистр строки, разбивая её на слова (максимальные alpha-раны)
    посимвольно. Не-буквенные символы (пробелы, пунктуация, цифры, \\t, \\n)
    копируются как есть и служат границами слов. Длина МОЖЕТ измениться только на
    редких символах, где ch.upper()/ch.lower() дают не один символ — это ловит
    хард-гейт длины на уровне сегмента у вызывающего кода."""
    out: list[str] = []
    i, n = 0, len(text)
    while i < n:
        if text[i].isalpha():
            j = i
            while j < n and text[j].isalpha():
                j += 1
            out.append(_normalize_word(text[i:j]))
            i = j
        else:
            out.append(text[i])
            i += 1
    return "".join(out)


# --- Разрешение эффективного all_caps / small_caps по цепочке наследования OOXML ---
#
# В OOXML all_caps/small_caps трёхзначны: True / False / None(«наследуй от родителя»).
# run.font.all_caps отражает ТОЛЬКО прямое форматирование рана (w:rPr/w:caps); капс,
# заданный СТИЛЕМ (частый в реальных договорах способ для шапок/преамбулы «СТОРОНЫ»),
# там None. Разрешение идёт по приоритету (высший → низший): прямой ран → стиль рана
# (character style) → стиль абзаца → базовые стили рекурсивно → стиль таблицы (для
# ячеек) → docDefaults. Первое НЕ-None значение выигрывает; явный False на любом
# уровне ОТМЕНЯЕТ унаследованный сверху True (не схлопываем трёхзначность в двузначность).
# Эмпирика уровней — см. отчёт по дофиксу наследования стиля.

_CAPS_TAG = {"all_caps": "w:caps", "small_caps": "w:smallCaps"}

# Сентинел «в кэше ничего нет» — отличает отсутствие ключа от валидного None
# (None здесь означает «наследуй», а не «не вычисляли»).
_MISSING = object()


def _onoff(element) -> bool | None:
    """Трёхзначное значение w:caps/w:smallCaps: None (элемента нет → наследуй),
    иначе булево по w:val (атрибут отсутствует → True; иначе '1'/'true'/'on' → True,
    '0'/'false'/'off' → False) — та же семантика, что у python-docx ST_OnOff."""
    if element is None:
        return None
    val = element.get(qn("w:val"))
    if val is None:
        return True
    return val in ("1", "true", "on")


def _read_docdefaults_caps(document, attr: str) -> bool | None:
    """docDefaults документа: w:docDefaults/w:rPrDefault/w:rPr/{w:caps|w:smallCaps}.
    python-docx это не выставляет через font-API, поэтому читаем XML напрямую (lxml
    уже в зависимостях). Возвращает True/False/None."""
    try:
        styles_el = document.styles.element
        dd = styles_el.find(qn("w:docDefaults"))
        if dd is None:
            return None
        rpr_default = dd.find(qn("w:rPrDefault"))
        if rpr_default is None:
            return None
        rpr = rpr_default.find(qn("w:rPr"))
        if rpr is None:
            return None
        return _onoff(rpr.find(qn(_CAPS_TAG[attr])))
    except Exception:
        # Битый/нестандартный styles.xml не должен ронять извлечение.
        return None


class _CapsResolver:
    """Разрешает эффективные all_caps/small_caps рана с учётом наследования стилей.

    Кэширует результат прохода по цепочке base_style ПО ОБЪЕКТУ СТИЛЯ (стилей в
    документе десятки, ранов — тысячи): без кэша каждый ран заново шёл бы вверх по
    цепочке. Защита от циклов (битые .docx с циклическим base_style существуют) —
    множество уже посещённых элементов стиля в пределах одного прохода.

    ЭТАП E'' — КЛЮЧ ТОЛЬКО ПО ОБЪЕКТУ, НИКОГДА ПО id(). Раньше и кеш, и защита от
    циклов ключевались ЧИСЛОМ `id(element)`, не удерживая ссылку на сам элемент.
    Это давало НЕДЕТЕРМИНИЗМ детекции: `paragraph.style`/`run.style` в python-docx —
    не кешированные свойства (новый прокси на каждое обращение), прокси lxml живёт,
    пока на него есть ссылка, а CPython ПЕРЕИСПОЛЬЗУЕТ освободившийся адрес. Поэтому
    другой стиль мог получить тот же id() и вернуть ЧУЖОЙ вердикт (ложное попадание),
    а в `seen` — сойти за уже посещённый (ложный цикл, обход обрывался и возвращал
    None). На реальном договоре (1097 стилей) это давало 1290 ложных попаданий на
    2396 вызовов и 4 разных detection_text на 4 попытки одного файла. Отсюда плыли
    и L2 (ORG/PER), и L3 (ADDRESS): detection_text лежит в metadata и читается
    per_search_view/anchor_search_view и detection_view, поэтому segment.text
    оставался побайтно стабильным, а детекция — нет.
    Ключ по САМОМУ объекту: dict/set удерживают ссылку → элемент не собирается →
    адрес не может быть переиспользован; lxml-элементы хешируются по идентичности,
    семантика прежняя. Память: несколько десятков элементов уже живого дерева."""

    def __init__(self, document) -> None:
        # ключ — (объект элемента стиля, attr); НЕ id(): см. docstring класса
        self._chain_cache: dict[tuple, bool | None] = {}
        self._docdefaults = {
            attr: _read_docdefaults_caps(document, attr)
            for attr in _CAPS_TAG
        }
        # ЭТАП O2: кэш «styleId -> объект стиля» НА ДОКУМЕНТ (см. style_of).
        self._style_by_id: dict[tuple, object] = {}

    def style_of(self, holder, style_type):
        """`paragraph.style` / `run.style` python-docx, но с кэшем по styleId.

        ЭТАП O2 (замер): геттер `.style` каждый раз идёт
        `part.get_style(styleId, type)` -> `styles.get_by_id` -> при styleId=None
        `styles.default(type)` -> `default_for(type)`, а тот ЛИНЕЙНО обходит ВСЕ
        стили документа, декодируя XML-атрибут типа у каждого. Получается
        O(абзацев x стилей): на фикстуре synthetic_corporate_large (5051 сегмент)
        это 12983 вызова и 18.5 c из 59 c пайплайна — 15% ТОЛЬКО на разрешение
        стиля по умолчанию, при том что ответ один и тот же.

        Результат `.style` зависит ТОЛЬКО от styleId элемента и от part документа
        (part у нас один на резолвер), поэтому кэш по (styleId, тип) — тождество,
        а не эвристика. Возвращается ТОТ ЖЕ объект-прокси вместо нового на каждое
        обращение; это безопасно (прокси только читают) и вдобавок стабилизирует
        ключи _chain_cache, которые и так были по `_element`.

        ВНИМАНИЕ (урок E''): ключ — СТРОКА styleId (или None), не id() объекта.
        Значение кэша удерживает сам прокси, поэтому переиспользования адреса,
        погубившего прошлую реализацию, здесь нет по построению."""
        element = holder._p if style_type is WD_STYLE_TYPE.PARAGRAPH else holder._element
        key = (element.style, style_type)
        style = self._style_by_id.get(key, _MISSING)
        if style is _MISSING:
            style = holder.style
            self._style_by_id[key] = style
        return style

    def _style_chain(self, style, attr: str) -> bool | None:
        """Первое НЕ-None значение attr вдоль цепочки base_style (сам стиль → его
        база → ...), с защитой от циклов и мемоизацией по стартовому стилю."""
        if style is None:
            return None
        key = (style._element, attr)   # объект, НЕ id() — см. docstring класса
        cached = self._chain_cache.get(key, _MISSING)
        if cached is not _MISSING:
            return cached

        result: bool | None = None
        seen = set()          # объекты элементов, НЕ id(): множество держит ссылку
        cur = style
        while cur is not None:
            el = cur._element
            if el in seen:   # циклическая цепочка стилей битого .docx
                break
            seen.add(el)
            value = getattr(cur.font, attr)   # прямое значение rPr этого стиля
            if value is not None:
                result = value
                break
            cur = cur.base_style

        self._chain_cache[key] = result
        return result

    def paragraph_baseline(self, paragraph, table_style) -> dict[str, bool | None]:
        """Разрешает уровни НИЖЕ рана (стиль абзаца → база → стиль таблицы → база →
        docDefaults) ОДИН РАЗ на абзац — они одинаковы для всех его ран. Возвращает
        {attr: True/False/None}. Вынесено из горячего пути: paragraph.style —
        дорогой геттер python-docx, звать его на каждый ран (×2 атрибута) незачем."""
        para_style = self.style_of(paragraph, WD_STYLE_TYPE.PARAGRAPH)
        baseline: dict[str, bool | None] = {}
        for attr in _CAPS_TAG:
            value = self._style_chain(para_style, attr)
            if value is None and table_style is not None:
                value = self._style_chain(table_style, attr)
            if value is None:
                value = self._docdefaults[attr]
            baseline[attr] = value
        return baseline

    def is_caps(self, run, baseline: dict[str, bool | None]) -> bool:
        """Итог для рана: эффективный all_caps ЛИБО small_caps == True. Приоритет:
        прямой ран → стиль рана (character style) → baseline абзаца-и-ниже. Явный
        False на верхнем уровне отменяет унаследованный True.

        run.style — дорогой геттер, поэтому берём его ЛЕНИВО и только если у рана
        реально задан character style (w:rStyle); у большинства ран его нет, и тогда
        уровень стиля рана пропускается вовсе."""
        rstyle = _MISSING  # ленивое разрешение run.style
        for attr in _CAPS_TAG:
            value = getattr(run.font, attr)   # 1. прямой ран
            if value is None:
                # 2. стиль рана — только если он вообще назначен
                if run._element.style is not None:
                    if rstyle is _MISSING:
                        rstyle = self.style_of(run, WD_STYLE_TYPE.CHARACTER)
                    value = self._style_chain(rstyle, attr)
                if value is None:
                    value = baseline[attr]    # 3. абзац-и-ниже (уже разрешён)
            if value is True:
                return True
        return False


def _paragraph_detection_text(paragraph: Paragraph, resolver: "_CapsResolver",
                              table_style=None) -> tuple[str, bool]:
    """Строит detection_text одного параграфа, нормализуя регистр в диапазонах
    ран с флагом all_caps/small_caps и копируя прочие раны как есть.

    Идём по iter_inner_content() (раны и гиперссылки в порядке документа) — именно
    из них python-docx 1.2.0 собирает paragraph.text, поэтому при отсутствии редких
    длина-меняющих символов результат посимвольно совпадает по позициям с
    paragraph.text. Возвращает (detection_text, есть_ли_маркированный_ран)."""
    # Уровни наследования ниже рана (стиль абзаца → база → стиль таблицы → docDefaults)
    # одинаковы для всех ран абзаца — разрешаем их один раз.
    baseline = resolver.paragraph_baseline(paragraph, table_style)
    parts: list[str] = []
    any_flag = False
    for item in paragraph.iter_inner_content():
        runs = item.runs if isinstance(item, Hyperlink) else (item,)
        for run in runs:
            flagged = resolver.is_caps(run, baseline)
            any_flag = any_flag or flagged
            parts.append(_normalize_case(run.text) if flagged else run.text)
    return "".join(parts), any_flag


def _maybe_set_detection_text(metadata: dict, text: str, detection_text: str, segment_id: str) -> None:
    """Хард-гейт равной длины: кладёт detection_text в metadata ТОЛЬКО если его
    длина совпадает с настоящим text. Иначе — не создаёт ключ вообще и логирует
    предупреждение (никогда не пытаемся «подогнать» соответствие)."""
    if len(detection_text) == len(text):
        metadata["detection_text"] = detection_text
    else:
        print(
            f"Предупреждение: detection_text не создан для сегмента {segment_id} — "
            f"нормализация изменила длину "
            f"({len(text)} -> {len(detection_text)}); текст: {text!r}",
            file=sys.stderr,
        )


def _cell_detection_text(cell, resolver: "_CapsResolver", table_style) -> tuple[str, bool]:
    """detection_text ячейки: по параграфам, склеенным '\\n' — ровно как
    _Cell.text = '\\n'.join(p.text for p in cell.paragraphs). table_style
    прокидывается вниз как самый низкий уровень наследования капса для ран ячейки."""
    parts: list[str] = []
    any_flag = False
    for p in cell.paragraphs:
        det, flag = _paragraph_detection_text(p, resolver, table_style)
        parts.append(det)
        any_flag = any_flag or flag
    return "\n".join(parts), any_flag


# --- Этап 2b: регистровая нормализация на уровне СЛОВА -----------------------
#
# Замер (baseline 26c37e3, полный корпус): из 1504 утёкших PER-сущностей 688
# (45.7%) стоят в тексте в аномальном регистре — recall на сплошь строчных ФИО
# РОВНО 0.0%, на ЗАГЛАВНЫХ 16.0%, на «пЕтРоВ» 53.0% (против 81.8% на нормальном
# регистре). Это потолок этапа.
#
# Ключевое различие, на котором построена вся конструкция: аномальность бывает
# СЛОВНАЯ и СЕГМЕНТНАЯ, и лечить их одинаково нельзя.
#
#   * ЗАГЛАВНОЕ слово и слово с ВНУТРЕННИМ переключением регистра ('пЕтРоВ')
#     аномальны САМИ ПО СЕБЕ, вне зависимости от соседей. Их можно чинить
#     точечно, и это дополнительно закрывает случай, мимо которого посегментная
#     эвристика проходила: 'Исполнитель: ПЕТРОВА ОЛЬГА ИВАНОВНА' — сегмент
#     смешанный, isupper() ложно, а ФИО внутри заглавное.
#
#   * СТРОЧНОЕ слово само по себе НЕ аномально — обычная русская проза почти
#     целиком из строчных слов. Приводить каждое строчное слово к Титульному =
#     превратить весь документ в Title Case и заставить Natasha «увидеть» имена
#     всюду. Поэтому строчный случай остаётся ПОСЕГМЕНТНЫМ: чиним, только если
#     сплошь строчный ВЕСЬ сегмент (у нормальной фразы есть хотя бы заглавная в
#     начале — сплошь строчный сегмент аномален как целое).
#
# Требование нумерации (_LIST_ITEM_RE) снято — это и была «известная непокрытая
# дыра» прошлой итерации. Длина сегмента и хард-гейт равной длины сохранены.
# Регистр меняется ТОЛЬКО в detection_text; segment.text не трогается, длина
# сохраняется посимвольно, поэтому оффсеты (и карта индексов нормализатора)
# остаются валидными в обоих текстах — спаны не съезжают.

_MIN_UPPER_WORD = 4  # слова 2-3 букв не трогаем: 'ООО'/'ИП'/'АО' — юр-формы, не капс

# Минимальная длина прогона заглавных слов, чтобы считать его ФИО, а не
# аббревиатурой/топонимом. Обоснование порога — в _caps_run_flags.
# ЭТАП C: возвращён к 2 (из 3). Порог 3 был вынужденной уступкой Stage2b-B: более
# точная маскировка ФИО обнажала адреса, случайно накрытые чужим NER-over-capture
# (двусловный заглавный прогон «ПЕТРОВА ОЛЬГА» не нормализовался, чтобы не открыть
# соседний адрес). После этапа C адрес — ЯКОРНЫЙ (Natasha-ADDRESS больше не хватает
# заглавные прогоны как топоним), поэтому обнажать нечего: порог 2 возвращает
# двусловные заглавные ФИО («ПЕТРОВА ОЛЬГА», «ИВАНОВ И.») в детекцию без роста утечки.
_MIN_CAPS_RUN = 2


def _is_invisible(ch: str) -> bool:
    """Невидимые управляющие форматирования (ZWSP/ZWNJ/ZWJ/WJ/SHY/BOM/LRM/RLM).

    Для разбиения на слова они ПРОЗРАЧНЫ — ровно так же, как в normalizer.py
    (проход 1 удаляет Cf целиком). Иначе 'о<SHY>льгА' распалось бы на 'о' и
    'льгА', одиночная 'о' не выглядела бы аномальной, и имя нормализовалось бы
    наполовину ('о Льга') — то есть омоглифная вставка внутрь слова отключала бы
    регистровую починку. Пробелоподобные (NBSP и пр.) сюда НЕ входят намеренно:
    нормализатор приводит их к обычному пробелу, т.е. считает границей слова —
    здесь мы обязаны считать так же."""
    return unicodedata.category(ch) == "Cf"


def _word_is_all_caps(word: str) -> bool:
    """Слово целиком заглавное и достаточно длинное, чтобы не быть юр-формой
    ('ООО'/'ИП'/'АО' — 2-3 буквы — не капс, а форма собственности)."""
    letters = [ch for ch in word if ch.isalpha()]
    return (len(letters) >= _MIN_UPPER_WORD
            and all(ch.isupper() for ch in letters))


def _word_has_inner_flip(word: str) -> bool:
    """Заглавная ПОСЛЕ строчной внутри одного слова ('пЕтРоВ'). В нормальном
    тексте так не пишут никогда, поэтому признак самодостаточен."""
    letters = [ch for ch in word if ch.isalpha()]
    if len(letters) < 2:
        return False
    seen_lower = False
    for ch in letters:
        if ch.islower():
            seen_lower = True
        elif ch.isupper() and seen_lower:
            return True
    return False


def _token_is_numeric_like(token: str) -> bool:
    """Токен (прогон непробельных) выглядит реквизитом: содержит ASCII-цифру.

    Регистр таких токенов ТРОГАТЬ НЕЛЬЗЯ. Свод цифровых омоглифов в normalizer.py
    (`_DIGIT_LOOKALIKE`) РЕГИСТРОЗАВИСИМ: 'б'->6 есть, 'Б'->6 НЕТ. Подняв регистр
    у '04Ч525ЗбО' (БИК, где буквы стоят вместо цифр), мы получаем 'ЗБО', свод
    цифр ломается, и реквизит перестаёт детектироваться ВООБЩЕ — то есть
    регистровая правка порождала бы утечку БИК/счёта. Замер это и показал:
    4 реквизита теряли маску полностью."""
    return any("0" <= ch <= "9" for ch in token)


def _is_mixed_alphabet(word: str) -> bool:
    """В слове есть и кириллица, и латиница — подпись омоглифной подмены букв
    ('иBаНоB', 'ПЧЕЛЫ ЕKAТЕРИНЫ'). Такими словами владеет свод `_ALPHA_FOLD`
    нормализатора, и он РЕГИСТРОЗАВИСИМ и неполон: 'B'->'В' есть, 'b'->'В' НЕТ.
    Опустив латинскую 'B' в 'b', мы лишаем нормализатор возможности свести её в
    кириллицу — слово остаётся смешанным, и ФИО не детектируется."""
    has_cyr = any("Ѐ" <= ch <= "ӿ" for ch in word)
    has_lat = any(("a" <= ch <= "z") or ("A" <= ch <= "Z") for ch in word)
    return has_cyr and has_lat


def _has_mixed_alphabet_word(text: str) -> bool:
    """Есть ли в сегменте хоть одно слово смешанного алфавита.

    Щит ставится на ВЕСЬ СЕГМЕНТ, а не на отдельное слово, и это принципиально.
    Замер словного щита (пропускать мимо только само омоглифное слово) показал
    результат ХУЖЕ обоих крайних вариантов: 7 ФИО терялись из-под маски ЦЕЛИКОМ,
    recall PER 67.2%->63.9%. Причина — ПОЛУнормализация: чистые слова имени
    приводятся к Титульному, омоглифное остаётся рваным, и Natasha перестаёт
    видеть имя целиком. Либо чиним регистр всему имени, либо не трогаем ничего.

    Поэтому сегмент с омоглифным словом этап 2b НЕ ТРОГАЕТ вовсе — стык двух
    нормализаций (регистр x омоглиф) вынесен отдельным шагом, см. FINDINGS
    Stage2b-A. Инвариант: 2b не имеет права создать НИ ОДНОЙ новой утечки."""
    return any(_is_mixed_alphabet(w) for w in re.split(r"\s+", text) if w)


def _numeric_token_ranges(text: str) -> list[tuple[int, int]]:
    """Диапазоны токенов (прогонов непробельных), похожих на реквизит.

    Именно ДИАПАЗОНЫ, а не «применить функцию потокенно»: словная ветвь смотрит
    на СОСЕДЕЙ слова (прогон заглавных, см. _caps_run_flags), и разрезание текста
    на токены до нормализации лишило бы её этого контекста — каждое слово стало
    бы одиноким, прогон никогда бы не находился."""
    ranges, i, n = [], 0, len(text)
    while i < n:
        if text[i].isspace():
            i += 1
            continue
        j = i
        while j < n and not text[j].isspace():
            j += 1
        if _token_is_numeric_like(text[i:j]):
            ranges.append((i, j))
        i = j
    return ranges


def _by_tokens(text: str, fn) -> str:
    """Применяет fn к тексту потокенно, пропуская мимо (байт-в-байт) токены,
    похожие на реквизит. Годится только для функций БЕЗ межсловного контекста
    (строчная ветвь); словная ветвь несёт щит внутри себя."""
    out, i, n = [], 0, len(text)
    while i < n:
        if text[i].isspace():
            out.append(text[i])
            i += 1
            continue
        j = i
        while j < n and not text[j].isspace():
            j += 1
        token = text[i:j]
        out.append(token if _token_is_numeric_like(token) else fn(token))
        i = j
    return "".join(out)


def _split_words(text: str) -> list[tuple[int, int]]:
    """Границы слов: прогоны букв, СКВОЗЬ невидимые (см. _is_invisible).
    Хвостовые невидимые в слово не забираем — они ничего не склеивают."""
    spans, i, n = [], 0, len(text)
    while i < n:
        if text[i].isalpha():
            j, last = i, i
            while j < n and (text[j].isalpha() or _is_invisible(text[j])):
                if text[j].isalpha():
                    last = j
                j += 1
            spans.append((i, last + 1))
            i = last + 1
        else:
            i += 1
    return spans


def _is_caps_like(word: str) -> bool:
    """Слово целиком заглавное (инициал 'Е' — тоже, это часть ФИО)."""
    letters = [ch for ch in word if ch.isalpha()]
    return bool(letters) and all(ch.isupper() for ch in letters)


def _caps_run_flags(text: str, spans: list[tuple[int, int]]) -> list[bool]:
    """Для каждого слова: входит ли оно в ПРОГОН из >= _MIN_CAPS_RUN подряд идущих
    заглавных слов (между ними — только пробелы/пунктуация, без цифр, зазор <= 3).

    Зачем прогон вообще: ОДИНОКОЕ заглавное слово среди текста нормального
    регистра — это АББРЕВИАТУРА ('код ОКВЭД', 'ОКТМО 45382000'), а не ФИО.
    Приводя её к Титульному, мы делаем из 'ОКВЭД' слово 'Оквэд', и Natasha
    читает его как фамилию: замер дал +17 ложных PER на негативах, 13 из них —
    ровно 'ОКВЭД'.

    Зачем порог именно 3: русское ФИО заглавными канонически трёхчастно —
    'ЕРШОВА ТАТЬЯНА ДМИТРИЕВНА', 'КРЫЛОВА Е. П.' (инициал — тоже часть прогона).
    А прогон из ДВУХ заглавных слов чаще оказывается топонимом без маркеров:
    замер поймал 'МОСКВА ЛЕНИНА 47/3 КВ 67' — этот адрес до 2b случайно попадал
    под NER-маску целиком (Natasha ловила заглавный прогон), а нормализация
    регистра эту случайную маску убирала, ОТКРЫВАЯ адрес. Порог 3 оставляет
    такие двусловные прогоны в покое. Цена — двусловные заглавные ФИО
    ('ПЕТРОВА ОЛЬГА', 'ИВАНОВ И.') 2b не чинит; это осознанный размен в пользу
    инварианта «ноль новых утечек реального текста»."""
    caps = [_is_caps_like(text[s:e]) for s, e in spans]
    flags = [False] * len(spans)
    k, n = 0, len(spans)
    while k < n:
        if not caps[k]:
            k += 1
            continue
        j = k
        while j + 1 < n and caps[j + 1]:
            gap = text[spans[j][1]:spans[j + 1][0]]
            if len(gap) > 3 or any(ch.isdigit() for ch in gap):
                break
            j += 1
        if j - k + 1 >= _MIN_CAPS_RUN:
            for t in range(k, j + 1):
                flags[t] = True
        k = j + 1
    return flags


def _max_word_run(text: str) -> int:
    """Длина самого длинного прогона подряд идущих СЛОВ (между ними — только
    пробелы/пунктуация, зазор <= 3 символов, без цифр).

    Нужна строчной ветви: она приводит к Титульному ВСЕ слова сплошь строчного
    сегмента, и на коротком обрывке это делает из месяца имя. Замер: '«25» мая
    2023 г.' -> '«25» Мая 2023 Г.', и 'Мая'/'Ноября' попадали в FP на негативах
    «дата документа — не ПДн» (5 штук, все на linebreak-мутациях, где перенос
    строки нарезает документ на короткие сегменты). У ФИО прогон слов всегда
    >= 3 ('кузнецовой марией дмитриевной'), а у обрывка даты — 1: между 'мая' и
    'г' стоит '2023', то есть цифры рвут прогон."""
    spans = _split_words(text)
    best = cur = 0
    for k, (s, e) in enumerate(spans):
        if k == 0:
            cur = 1
        else:
            gap = text[spans[k - 1][1]:s]
            cur = cur + 1 if (len(gap) <= 3 and not any(c.isdigit() for c in gap)) else 1
        best = max(best, cur)
    return best


def _normalize_anomalous_words(text: str) -> str:
    """Нормализует регистр ТОЛЬКО у аномальных слов; остальные слова и все
    не-буквенные символы копируются байт-в-байт. Тем самым текст с нормальным
    регистром проходит через функцию тождественно — это и есть гарантия, что уже
    работающая детекция не заденется.

    Две ветви аномалии разведены по строгости:
      * ВНУТРЕННИЙ флип ('пЕтРоВ') — так не пишут никогда, чиним всегда;
      * ЗАГЛАВНОЕ слово — чиним, только если оно в прогоне заглавных (см.
        _caps_run_flags), иначе это аббревиатура, а не ФИО."""
    spans = _split_words(text)
    runs = _caps_run_flags(text, spans)
    numeric = _numeric_token_ranges(text)
    out, prev = [], 0
    for (s, e), in_run in zip(spans, runs):
        out.append(text[prev:s])
        word = text[s:e]
        in_numeric = any(ts <= s and e <= te for ts, te in numeric)
        fix = (not in_numeric) and (
            _word_has_inner_flip(word) or (in_run and _word_is_all_caps(word)))
        out.append(_normalize_word(word) if fix else word)
        prev = e
    out.append(text[prev:])
    return "".join(out)


def _apply_lowercase_heuristic(doc: SourceDocument) -> None:
    """Изменение 2 + этап 2b: эвристика для НАСТОЯЩЕГО строчного/заглавного/
    рваного ввода без форматирования (одинаково для .docx и .txt). Применяется к
    сегментам, у которых detection_text ещё НЕ установлен Изменением 1, и только
    к коротким (< _MAX_HEURISTIC_LEN). Две ветви, см. блок комментариев выше:

      * сплошь строчный сегмент  -> нормализуем сегмент целиком (_normalize_case);
      * иначе                    -> нормализуем ТОЛЬКО аномальные слова
                                    (заглавные / с внутренним флипом регистра).

    Тот же посимвольный механизм и тот же хард-гейт равной длины, что в
    Изменении 1: если длина не сохранилась, detection_text не создаётся вовсе."""
    # Щит combo — на ВЕСЬ ДОКУМЕНТ, а не на отдельный сегмент. Посегментный щит
    # замерен и оказался дырявым: сегмент с омоглифами мы пропускали, но соседние
    # чинили, PER-маски в документе смещались, а вслед за ними жадный
    # `yargy.AddrExtractor` (находка 0c-C/D) перестраивал свои спаны — и адрес в
    # ДРУГОМ сегменте того же документа открывался. Так 2b создавал новые утечки
    # реального текста на combo-документах, хотя «не трогал» их сегменты.
    #
    # На combo-документе откатываемся РОВНО к историческому поведению (сплошь
    # одно-регистровый сегмент + пункт перечисления), а НЕ к бездействию: полное
    # отключение отняло бы то, что работало до 2b, и само роняло детекцию адреса.
    # 2b на combo обязан не добавить ничего нового и не отнять ничего старого.
    combo = any(_has_mixed_alphabet_word(s.text) for s in doc.segments if s.text)
    for seg in doc.segments:
        text = seg.text
        if not text or "detection_text" in seg.metadata:
            continue
        if len(text) > _MAX_HEURISTIC_LEN:
            continue
        if combo:
            if (text.islower() or text.isupper()) and _LIST_ITEM_RE.match(text):
                _maybe_set_detection_text(seg.metadata, text,
                                          _normalize_case(text), seg.id)
            continue
        if text.islower():
            # Строчная ветвь правит ВСЕ слова сегмента, поэтому требует признака
            # имени — прогон слов >= _MIN_CAPS_RUN (см. _max_word_run), иначе
            # короткий обрывок вроде '«25» мая 2023 г.' даёт ложное 'Мая'.
            if _max_word_run(text) < _MIN_CAPS_RUN:
                continue
            candidate = _by_tokens(text, _normalize_case)
        else:
            candidate = _normalize_anomalous_words(text)   # щит внутри
        if candidate == text:
            continue  # регистр в порядке — detection_text не нужен
        _maybe_set_detection_text(seg.metadata, text, candidate, seg.id)


def _extract_docx(path: str) -> SourceDocument:
    try:
        document = Document(path)
    except (PackageNotFoundError, zipfile.BadZipFile) as exc:
        # Файл не является валидным .docx-контейнером: переименованный .txt,
        # обрезанный/битый или пустой файл. python-docx кидает PackageNotFoundError
        # (или BadZipFile) — оборачиваем в понятный ValueError, который CLI ловит.
        raise ValueError(
            f"Файл не является корректным .docx (повреждён или неверный формат): {path}"
        ) from exc
    segments: list[TextSegment] = []

    # Один резолвер капса на документ: держит кэш проходов по цепочкам стилей и
    # прочитанные умолчания docDefaults (см. _CapsResolver).
    caps_resolver = _CapsResolver(document)

    paragraph_counter = 0
    table_counter = 0
    seen_tcs = set()

    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            paragraph = Paragraph(child, document)
            # ЭТАП O2: через кэш резолвера — раньше это были ДВА полных разрешения
            # стиля на абзац (проверка на None и .name), плюс третье в
            # paragraph_baseline. См. _CapsResolver.style_of.
            _p_style = caps_resolver.style_of(paragraph, WD_STYLE_TYPE.PARAGRAPH)
            style_name = _p_style.name if _p_style is not None else None
            seg_id = f"p{paragraph_counter}"
            metadata = {"paragraph_index": paragraph_counter, "style": style_name}
            # Изменение 1: detection_text из форматирования all_caps/small_caps
            # (прямого или унаследованного от стиля абзаца/символа).
            det_text, any_flag = _paragraph_detection_text(paragraph, caps_resolver)
            if any_flag:
                _maybe_set_detection_text(metadata, paragraph.text, det_text, seg_id)
            segments.append(
                TextSegment(
                    id=seg_id,
                    text=paragraph.text,
                    source_type="docx_paragraph",
                    metadata=metadata,
                )
            )
            paragraph_counter += 1
        elif child.tag == qn("w:tbl"):
            table = Table(child, document)
            table_idx = table_counter
            # Стиль таблицы — самый низкий уровень наследования капса для ран ячеек
            # (python-docx не отражает его в run/paragraph style). Берём один раз.
            table_style = table.style
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    # Этап 1b: вложенная таблица здесь БОЛЬШЕ НЕ логируется. Раньше
                    # на неё печаталось предупреждение в stderr — и этим всё
                    # заканчивалось: текст молча выпадал из результата, а warning
                    # терялся среди прочего вывода. Теперь непрочитанные зоны
                    # (включая вложенные таблицы) обнаруживает unread_zones.
                    # scan_unread_zones, а решение принимает политика блока 7:
                    # по умолчанию encrypt ОТКАЗЫВАЕТСЯ, с --allow-lossy — пишет
                    # текст зон в {sid}.unread.json. Дублировать здесь warning
                    # незачем: extract() — библиотечная функция, её задача извлечь
                    # то, что она умеет, а не решать за вызывающего.
                    tc = cell._tc
                    cell_seg_id = f"t{table_idx}_r{row_idx}_c{col_idx}"
                    metadata = {
                        "table_index": table_idx,
                        "row_index": row_idx,
                        "col_index": col_idx,
                    }
                    if tc in seen_tcs:
                        cell_text = ""
                    else:
                        seen_tcs.add(tc)
                        cell_text = cell.text
                        # Изменение 1: detection_text ячейки из форматирования ран
                        # (прямого или унаследованного от стиля ячейки/таблицы).
                        det_text, any_flag = _cell_detection_text(cell, caps_resolver, table_style)
                        if any_flag:
                            _maybe_set_detection_text(metadata, cell_text, det_text, cell_seg_id)
                    segments.append(
                        TextSegment(
                            id=cell_seg_id,
                            text=cell_text,
                            source_type="docx_table_cell",
                            metadata=metadata,
                        )
                    )
            table_counter += 1

    return SourceDocument(segments=segments, source_format="docx", source_path=path)


def _looks_like_mojibake(text: str) -> bool:
    """Эвристика тихой порчи: высокая доля управляющих/replacement-символов.

    Байты UTF-16 (особенно UTF-16-LE, где ASCII-символ = <буква>0x00) успешно
    «декодируются» как utf-8 или cp1251 БЕЗ исключения — но результат забит
    NUL'ами и другими C0/C1-управляющими символами, которых в нормальном тексте
    договора нет. Если их доля велика — декодирование заведомо мусорное. Обычный
    UTF-8/cp1251-текст (в т.ч. с эмодзи и переводами строк) даёт долю ~0, поэтому
    порог 0.30 не задевает валидные файлы, но уверенно ловит UTF-16-моджибейк
    (у него доля ~0.5). Не пытаемся УГАДАТЬ UTF-16 без BOM — только не даём
    испорченному тексту тихо пройти как валидный."""
    if not text:
        return False
    bad = 0
    for ch in text:
        code = ord(ch)
        if ch == "�":
            bad += 1  # replacement-символ: часть байтов не декодировалась
        elif code < 0x20 and ch not in "\t\n\r":
            bad += 1  # C0-управляющие (в т.ч. NUL из UTF-16), кроме обычных пробелов
        elif 0x7f <= code <= 0x9f:
            bad += 1  # DEL и C1-управляющие
    return bad / len(text) > 0.30


def _extract_txt(path: str) -> SourceDocument:
    with open(path, "rb") as f:
        data = f.read()

    # UTF-16 с BOM (LE 0xFF 0xFE / BE 0xFE 0xFF) — так блокнот Windows сохраняет
    # "Юникод". Детектируем ДО отката на cp1251: иначе utf-8-sig падает, cp1251
    # молча декодирует UTF-16-байты в моджибейк, и ПДн (ИНН/телефоны/ФИО) не
    # находятся детекторами — тихая порча текста и утечка ПДн в искажённом виде.
    if data.startswith(b"\xff\xfe") or data.startswith(b"\xfe\xff"):
        encoding = "utf-16"
        raw = data.decode(encoding)
    else:
        encoding = "utf-8-sig"
        try:
            raw = data.decode(encoding)
        except UnicodeDecodeError:
            encoding = "cp1251"
            raw = data.decode(encoding)

    # B5-fix: UTF-16 БЕЗ BOM «декодируется» как utf-8-sig/cp1251 без исключения, но
    # в моджибейк — ПДн не находятся и утекают искажёнными без единого предупреждения
    # (тот же класс тихой порчи, что закрывал фикс #7, но #7 ловит только UTF-16 C BOM).
    # Надёжного автоопределения кодировки без BOM не существует, поэтому не угадываем:
    # если ни один из перепробованных вариантов не дал «чистого» текста — поднимаем
    # явную ошибку. Лучше внятный отказ, чем тихая утечка ПДн в LLM.
    if _looks_like_mojibake(raw):
        raise ValueError(
            f"Не удалось надёжно определить кодировку файла: {path}. "
            f"Сохраните файл в UTF-8 и повторите"
        )

    text = raw.replace("\r\n", "\n").replace("\r", "\n")
    lines = text.split("\n")
    if lines and lines[-1] == "":
        lines.pop()

    segments = [
        TextSegment(
            id=f"l{index}",
            text=line,
            source_type="txt_line",
            metadata={"line_index": index, "encoding": encoding},
        )
        for index, line in enumerate(lines)
    ]

    return SourceDocument(segments=segments, source_format="txt", source_path=path)


def extract(path: str) -> SourceDocument:
    if not Path(path).is_file():
        raise FileNotFoundError(path)

    ext = Path(path).suffix.lower()
    if ext == ".docx":
        doc = _extract_docx(path)
    elif ext == ".txt":
        doc = _extract_txt(path)
    else:
        raise ValueError(f"Неподдерживаемый формат: {ext}. Поддерживаются: .docx, .txt")

    # Изменение 2: узкая эвристика для настоящего строчного/заглавного ввода —
    # для .docx и .txt, поверх сегментов, у которых detection_text не задан.
    _apply_lowercase_heuristic(doc)
    return doc
