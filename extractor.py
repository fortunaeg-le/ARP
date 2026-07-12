import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

from models import SourceDocument, TextSegment


def _extract_docx(path: str) -> SourceDocument:
    try:
        document = Document(path)
    except (PackageNotFoundError, zipfile.BadZipFile) as exc:
        # Файл не является валидным .docx-контейнером: переименованный .txt,
        # обрезанный/битый или пустой файл. python-docx кидает PackageNotFoundError
        # (или BadZipFile) — оборачиваем в понятный ValueError, который CLI ловит.
        raise ValueError(
            f"Файл не является корректным .docx (повреждён или неверный формат): {path}"
        ) from exc
    segments: list[TextSegment] = []

    paragraph_counter = 0
    table_counter = 0
    seen_tcs = set()

    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            paragraph = Paragraph(child, document)
            style_name = paragraph.style.name if paragraph.style is not None else None
            segments.append(
                TextSegment(
                    id=f"p{paragraph_counter}",
                    text=paragraph.text,
                    source_type="docx_paragraph",
                    metadata={"paragraph_index": paragraph_counter, "style": style_name},
                )
            )
            paragraph_counter += 1
        elif child.tag == qn("w:tbl"):
            table = Table(child, document)
            table_idx = table_counter
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    if len(cell.tables) > 0:
                        print(
                            f"Предупреждение: вложенная таблица в ячейке "
                            f"t{table_idx}_r{row_idx}_c{col_idx} не извлечена (MVP)",
                            file=sys.stderr,
                        )
                    tc = cell._tc
                    if tc in seen_tcs:
                        cell_text = ""
                    else:
                        seen_tcs.add(tc)
                        cell_text = cell.text
                    segments.append(
                        TextSegment(
                            id=f"t{table_idx}_r{row_idx}_c{col_idx}",
                            text=cell_text,
                            source_type="docx_table_cell",
                            metadata={
                                "table_index": table_idx,
                                "row_index": row_idx,
                                "col_index": col_idx,
                            },
                        )
                    )
            table_counter += 1

    return SourceDocument(segments=segments, source_format="docx", source_path=path)


def _looks_like_mojibake(text: str) -> bool:
    """Эвристика тихой порчи: высокая доля управляющих/replacement-символов.

    Байты UTF-16 (особенно UTF-16-LE, где ASCII-символ = <буква>0x00) успешно
    «декодируются» как utf-8 или cp1251 БЕЗ исключения — но результат забит
    NUL'ами и другими C0/C1-управляющими символами, которых в нормальном тексте
    договора нет. Если их доля велика — декодирование заведомо мусорное. Обычный
    UTF-8/cp1251-текст (в т.ч. с эмодзи и переводами строк) даёт долю ~0, поэтому
    порог 0.30 не задевает валидные файлы, но уверенно ловит UTF-16-моджибейк
    (у него доля ~0.5). Не пытаемся УГАДАТЬ UTF-16 без BOM — только не даём
    испорченному тексту тихо пройти как валидный."""
    if not text:
        return False
    bad = 0
    for ch in text:
        code = ord(ch)
        if ch == "�":
            bad += 1  # replacement-символ: часть байтов не декодировалась
        elif code < 0x20 and ch not in "\t\n\r":
            bad += 1  # C0-управляющие (в т.ч. NUL из UTF-16), кроме обычных пробелов
        elif 0x7f <= code <= 0x9f:
            bad += 1  # DEL и C1-управляющие
    return bad / len(text) > 0.30


def _extract_txt(path: str) -> SourceDocument:
    with open(path, "rb") as f:
        data = f.read()

    # UTF-16 с BOM (LE 0xFF 0xFE / BE 0xFE 0xFF) — так блокнот Windows сохраняет
    # "Юникод". Детектируем ДО отката на cp1251: иначе utf-8-sig падает, cp1251
    # молча декодирует UTF-16-байты в моджибейк, и ПДн (ИНН/телефоны/ФИО) не
    # находятся детекторами — тихая порча текста и утечка ПДн в искажённом виде.
    if data.startswith(b"\xff\xfe") or data.startswith(b"\xfe\xff"):
        encoding = "utf-16"
        raw = data.decode(encoding)
    else:
        encoding = "utf-8-sig"
        try:
            raw = data.decode(encoding)
        except UnicodeDecodeError:
            encoding = "cp1251"
            raw = data.decode(encoding)

    # B5-fix: UTF-16 БЕЗ BOM «декодируется» как utf-8-sig/cp1251 без исключения, но
    # в моджибейк — ПДн не находятся и утекают искажёнными без единого предупреждения
    # (тот же класс тихой порчи, что закрывал фикс #7, но #7 ловит только UTF-16 C BOM).
    # Надёжного автоопределения кодировки без BOM не существует, поэтому не угадываем:
    # если ни один из перепробованных вариантов не дал «чистого» текста — поднимаем
    # явную ошибку. Лучше внятный отказ, чем тихая утечка ПДн в LLM.
    if _looks_like_mojibake(raw):
        raise ValueError(
            f"Не удалось надёжно определить кодировку файла: {path}. "
            f"Сохраните файл в UTF-8 и повторите"
        )

    text = raw.replace("\r\n", "\n").replace("\r", "\n")
    lines = text.split("\n")
    if lines and lines[-1] == "":
        lines.pop()

    segments = [
        TextSegment(
            id=f"l{index}",
            text=line,
            source_type="txt_line",
            metadata={"line_index": index, "encoding": encoding},
        )
        for index, line in enumerate(lines)
    ]

    return SourceDocument(segments=segments, source_format="txt", source_path=path)


def extract(path: str) -> SourceDocument:
    if not Path(path).is_file():
        raise FileNotFoundError(path)

    ext = Path(path).suffix.lower()
    if ext == ".docx":
        return _extract_docx(path)
    elif ext == ".txt":
        return _extract_txt(path)
    else:
        raise ValueError(f"Неподдерживаемый формат: {ext}. Поддерживаются: .docx, .txt")
