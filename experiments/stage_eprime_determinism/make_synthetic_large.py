# -*- coding: utf-8 -*-
"""
Синтетический двойник реального договора для этапа E'' (детерминизм детекции).

ПОЛНОСТЬЮ ВЫДУМАННЫЕ данные (ни одного реального ФИО/адреса/реквизита). Структура
скопирована из наблюдений по реальному документу (см. память сессии, обезличенно):
  - объём ~2500 сегментов (абзацы + ячейки таблиц) -- размер критичен для репро;
  - раздел определений: '"X" означает ...', 'под "X" понимается ...',
    'именуемое в дальнейшем "X"' -- класс, что до этапа C' давал тысячи ORG-масок;
  - ссылки на законы с аббревиатурой внутри ('Закона об АО', 'в соответствии с ФЗ');
  - коллективы-роли во множественном числе в разных падежах (акционеры/участники/стороны);
  - реквизитные таблицы: несколько организаций, юр./факт. адрес, ИНН/ОГРН/КПП/БИК/счёт;
  - ФИО в подписях с инициалами и в косвенных падежах;
  - переносы строк внутри значений (адрес/счёт/название) -- multispan-путь E/E'.

Контрольные суммы ИНН/ОГРН -- ВАЛИДНЫЕ (тот же алгоритм, что src/regex_detector.py),
иначе якорная детекция может вести себя иначе, чем на реальных реквизитах.
"""
import os
import random

import docx
from docx.shared import Pt

ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
OUT_DOCX = os.path.join(ROOT, "tests", "fixtures", "synthetic_corporate_large.docx")
OUT_TXT = os.path.join(ROOT, "tests", "fixtures", "synthetic_corporate_large.txt")

rnd = random.Random(20260723)


def valid_inn10(seed_digits):
    d = seed_digits[:9]
    weights = [2, 4, 10, 3, 5, 9, 4, 6, 8]
    c = sum(w * x for w, x in zip(weights, d)) % 11 % 10
    return "".join(map(str, d)) + str(c)


def valid_inn12(seed_digits):
    d = seed_digits[:10]
    w11 = [7, 2, 4, 10, 3, 5, 9, 4, 6, 8]
    w12 = [3, 7, 2, 4, 10, 3, 5, 9, 4, 6, 8]
    c11 = sum(w * x for w, x in zip(w11, d)) % 11 % 10
    d11 = d + [c11]
    c12 = sum(w * x for w, x in zip(w12, d11)) % 11 % 10
    return "".join(map(str, d)) + str(c11) + str(c12)


def valid_ogrn13(seed_digits):
    d = seed_digits[:12]
    n = int("".join(map(str, d)))
    c = n % 11 % 10
    return "".join(map(str, d)) + str(c)


def rand_digits(n):
    return [rnd.randint(0, 9) for _ in range(n)]


def gen_inn10():
    return valid_inn10(rand_digits(9))


def gen_inn12():
    return valid_inn12(rand_digits(10))


def gen_ogrn13():
    return valid_ogrn13(rand_digits(12))


def gen_kpp():
    return "".join(str(rnd.randint(0, 9)) for _ in range(9))


def gen_bik():
    return "04" + "".join(str(rnd.randint(0, 9)) for _ in range(7))


def gen_account():
    return "".join(str(rnd.randint(0, 9)) for _ in range(20))


def gen_snils():
    return "%03d-%03d-%03d %02d" % (
        rnd.randint(100, 999), rnd.randint(100, 999), rnd.randint(100, 999),
        rnd.randint(10, 99))


FIRST_NAMES_M = ["Александр", "Дмитрий", "Сергей", "Андрей", "Максим", "Игорь", "Николай"]
PATRONYMICS_M = ["Александрович", "Дмитриевич", "Сергеевич", "Андреевич", "Игоревич"]
SURNAMES_M = ["Воронцов", "Матвеенко", "Гриневич", "Соломин", "Черкасов", "Дубровин", "Ясенев"]
FIRST_NAMES_F = ["Елена", "Мария", "Ольга", "Наталья", "Анна", "Виктория"]
PATRONYMICS_F = ["Александровна", "Дмитриевна", "Сергеевна", "Андреевна", "Игоревна"]
SURNAMES_F = ["Воронцова", "Матвеенко", "Гриневич", "Соломина", "Черкасова", "Дубровина"]

ORG_CORES = ["Меридиан", "Кварцит", "Спектрум-Инвест", "Вектор Строй", "Гелиос Трейд",
             "Апекс Логистика", "Родник-Финанс", "Технопарк Юг", "Орбита-Сервис"]
ORG_FORMS = ["Общество с ограниченной ответственностью", "Акционерное общество",
             "Публичное акционерное общество"]
ORG_ABBR = {"Общество с ограниченной ответственностью": "ООО",
            "Акционерное общество": "АО",
            "Публичное акционерное общество": "ПАО"}

CITIES = ["Новоуральск", "Западнодвинск", "Краснокамск", "Южноуральск", "Березники"]
STREETS = ["ул. Полевая", "пр-т Мирный", "ул. Заречная", "пер. Садовый", "ш. Промышленное"]

ROLE_TERMS_PLURAL = [
    ("акционеры", "акционеров", "акционерам"),
    ("участники", "участников", "участникам"),
    ("стороны", "сторон", "сторонам"),
    ("работники", "работников", "работникам"),
]

LAW_REFS = [
    "в соответствии с Федеральным законом «Об акционерных обществах» (далее -- Закон об АО)",
    "согласно ст. 65.2 ГК РФ и Закону об ООО",
    "в порядке, предусмотренном Законом о банкротстве",
    "с учётом требований ФЗ «О персональных данных»",
]


def full_name(male=True):
    if male:
        return "%s %s %s" % (rnd.choice(SURNAMES_M), rnd.choice(FIRST_NAMES_M), rnd.choice(PATRONYMICS_M))
    return "%s %s %s" % (rnd.choice(SURNAMES_F), rnd.choice(FIRST_NAMES_F), rnd.choice(PATRONYMICS_F))


def initials(fio):
    parts = fio.split()
    return "%s %s.%s." % (parts[0], parts[1][0], parts[2][0])


def oblique_name_dative(male=True):
    # косвенный падеж (дательный) -- намеренно рвём на 2 слова структурой предложения,
    # как это делает Natasha на реальных договорах (Stage B PER-B класс)
    if male:
        s, f, p = rnd.choice(SURNAMES_M), rnd.choice(FIRST_NAMES_M), rnd.choice(PATRONYMICS_M)
        return "%sу %sу %sу" % (s, f[:-1], p[:-2])
    s, f, p = rnd.choice(SURNAMES_F), rnd.choice(FIRST_NAMES_F), rnd.choice(PATRONYMICS_F)
    return "%sой %s %sе" % (s, f[:-1] + "и", p[:-1] + "е")


def org_name():
    form = rnd.choice(ORG_FORMS)
    core = rnd.choice(ORG_CORES)
    abbr = ORG_ABBR[form]
    return form, abbr, core


def definitions_paragraph(idx):
    """Раздел определений -- класс, что до C' давал тысячи ORG-масок."""
    term = rnd.choice(["Договор", "Стороны", "Объект", "Услуги", "Продукция",
                        "Конфиденциальная информация", "Форс-мажор", "Отчётный период"])
    style = idx % 3
    if style == 0:
        return "«%s» означает документ, регулирующий отношения Сторон по пункту %d настоящего договора." % (term, idx)
    if style == 1:
        return "Под «%s» понимается совокупность условий, изложенных в разделе %d." % (term, idx % 20 + 1)
    return "именуемое в дальнейшем «%s» для целей раздела %d." % (term, idx % 20 + 1)


def law_paragraph(idx):
    return "Стороны действуют %s (пункт %d)." % (rnd.choice(LAW_REFS), idx)


def role_paragraph(idx):
    nom, gen, dat = rnd.choice(ROLE_TERMS_PLURAL)
    forms = [
        "Все %s несут солидарную ответственность по обязательствам." % nom,
        "Настоящее положение распространяется на всех %s без исключения." % gen,
        "Уведомление направляется всем %s в течение 10 рабочих дней." % dat,
    ]
    return forms[idx % 3]


def signature_block(male):
    fio = full_name(male)
    ini = initials(fio)
    return "Подпись: _____________ / %s (%s) /" % (ini, fio)


def requisite_table(document, n_orgs):
    for i in range(n_orgs):
        table = document.add_table(rows=8, cols=2)
        form, abbr, core = org_name()
        inn = gen_inn10() if i % 2 == 0 else gen_inn12()
        ogrn = gen_ogrn13()
        kpp = gen_kpp()
        bik = gen_bik()
        acc = gen_account()
        city = rnd.choice(CITIES)
        street = rnd.choice(STREETS)
        house = rnd.randint(1, 200)
        rows_data = [
            ("Наименование", "%s «%s»" % (form, core)),
            ("Юридический адрес",
             "%d%03d, г. %s,\n%s, д. %d" % (rnd.randint(1, 6), rnd.randint(0, 999), city, street, house)),
            ("Фактический адрес",
             "%d%03d, г. %s,\n%s, д. %d, оф. %d" % (
                 rnd.randint(1, 6), rnd.randint(0, 999), rnd.choice(CITIES), rnd.choice(STREETS), house, i + 1)),
            ("ИНН", inn),
            ("ОГРН", ogrn),
            ("КПП", kpp),
            ("Банковские реквизиты",
             "р/с %s\nБИК %s" % (acc, bik)),
            ("Подписант", "%s, действующий на основании Устава" % full_name(i % 2 == 0)),
        ]
        for r, (label, value) in enumerate(rows_data):
            table.cell(r, 0).text = label
            table.cell(r, 1).text = value
        document.add_paragraph("")


def build():
    document = docx.Document()
    style = document.styles["Normal"]
    style.font.size = Pt(11)

    document.add_paragraph("ДОГОВОР № КРУПНЫЙ-2026/СИНТ").alignment = 1

    # --- преамбула с ФИО и ролями-коллективами (person anchors + role plural) ---
    for i in range(60):
        p = full_name(male=(i % 2 == 0))
        document.add_paragraph(
            "%s, именуемый в дальнейшем «%s», действующий от имени %s," % (
                p, rnd.choice(["Сторона %d" % i, "Представитель"]), oblique_name_dative(i % 2 == 0)))

    # --- крупный раздел определений: ~1800 абзацев ---
    for i in range(1800):
        document.add_paragraph(definitions_paragraph(i))

    # --- ссылки на законы: ~600 абзацев ---
    for i in range(600):
        document.add_paragraph(law_paragraph(i))

    # --- роли-коллективы во множественном числе: ~600 абзацев ---
    for i in range(600):
        document.add_paragraph(role_paragraph(i))

    # --- реквизитные таблицы: 70 организаций x 8 строк x 2 столбца = 1120 сегментов ---
    requisite_table(document, 70)

    # --- ФИО в подписях, инициалы, косвенный падеж: ~600 абзацев ---
    for i in range(300):
        document.add_paragraph(signature_block(i % 2 == 0))
        document.add_paragraph("СНИЛС %s подписанта учтён в кадровом деле." % gen_snils())

    # --- хвост: ещё немного определений ---
    for i in range(1800, 2000):
        document.add_paragraph(definitions_paragraph(i))

    os.makedirs(os.path.dirname(OUT_DOCX), exist_ok=True)
    document.save(OUT_DOCX)

    # txt-двойник (для проверки: воспроизводится ли эффект вне .docx-пути извлечения)
    lines = []
    for p in document.paragraphs:
        lines.append(p.text)
    with open(OUT_TXT, "w", encoding="utf-8") as f:
        f.write("\n".join(lines) + "\n")

    print("saved:", OUT_DOCX)
    print("saved:", OUT_TXT)


if __name__ == "__main__":
    build()
