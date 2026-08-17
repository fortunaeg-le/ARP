# -*- coding: utf-8 -*-
"""
E'' — ДОКАЗАТЕЛЬСТВО МЕХАНИЗМА: id()-ключ в _CapsResolver (src/extractor.py).

Подозреваемое место (src/extractor.py, _CapsResolver._style_chain):
    key = (id(style._element), attr)      # <- КЛЮЧ = ЧИСЛО id(), ССЫЛКА НЕ ДЕРЖИТСЯ
    ...
    el_id = id(cur._element)              # <- то же в защите от циклов
    seen.add(el_id)

Почему это даёт недетерминизм:
  * `paragraph.style` / `run.style` в python-docx — НЕ кешированные свойства: на
    каждое обращение создаётся НОВЫЙ объект стиля, а lxml создаёт прокси элемента
    по требованию и освобождает, когда исчезает последняя ссылка;
  * кеш хранит только ЧИСЛО id(...), ссылку на объект не удерживает -> объект
    уничтожается сборщиком мусора;
  * CPython ПЕРЕИСПОЛЬЗУЕТ освободившиеся адреса -> ДРУГОЙ стиль может получить
    ТОТ ЖЕ id() -> ЛОЖНОЕ ПОПАДАНИЕ: возвращается результат ДРУГОГО стиля;
  * то же в `seen`: прежний `cur` освобождается, следующий может занять его адрес
    -> ЛОЖНЫЙ ЦИКЛ -> обход цепычки обрывается и возвращает None вместо значения.
Момент сборки мусора и поведение аллокатора между прогонами не совпадают =>
результат МЕНЯЕТСЯ В ОДНОМ ПРОЦЕССЕ при идентичном входе.

Следствие для детекции: _CapsResolver -> all_caps/small_caps -> detection_text
(лежит в metadata, НЕ в segment.text!) -> его читают per_search_view/
anchor_search_view (L2, ORG/PER) и detection_view (L3, ADDRESS). Поэтому
segment.text побайтно стабилен, а L2 и L3 плавают ОБА — ровно картина владельца.

ЗАПУСК (лучше всего на РЕАЛЬНОМ .docx из Word — там десятки стилей):
    venv\\Scripts\\python.exe experiments\\stage_eprime_determinism\\probe_capsresolver_id.py <файл.docx>
Вывод — только служебный (имена СТИЛЕЙ, не текст документа).
"""
import os
import sys
import gc

ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
sys.path.insert(0, os.path.join(ROOT, "src"))
os.chdir(ROOT)

from docx import Document
import extractor as EX

DOC = sys.argv[1] if len(sys.argv) > 1 else os.path.join(
    ROOT, "tests", "fixtures", "synthetic_corporate_large.docx")
if not os.path.isfile(DOC) or not DOC.lower().endswith(".docx"):
    cand = os.path.join(ROOT, "tests", "corpus", "docs")
    DOC = os.path.join(cand, sorted(
        f for f in os.listdir(cand) if f.lower().endswith(".docx"))[0])
print("документ:", os.path.basename(DOC))
document = Document(DOC)
print("стилей в документе:", len(list(document.styles)), "\n")

print("=== ТЕСТ 1: python-docx создаёт НОВЫЙ объект стиля на каждое обращение? ===")
target = None
for p in document.paragraphs:
    if p.style is not None:
        target = p
        break
if target is None:
    print("  (в document.paragraphs нет абзацев со стилем — пропуск)")
else:
    print("  p.style is p.style           :", target.style is target.style,
          "  <- False означает: объект КАЖДЫЙ РАЗ НОВЫЙ")
    held = target.style._element
    print("  id стабилен, пока ДЕРЖИМ ссылку:", id(held) == id(target.style._element))
    del held

print("\n=== ТЕСТ 2: ПРЯМАЯ коллизия — освободить стиль A, занять адрес стилем B ===")
names = [s.name for s in document.styles if s.name]
hits = 0
if len(names) >= 2:
    for a_name, b_name in zip(names, names[1:]):
        try:
            ea = document.styles[a_name]._element
            ia = id(ea)
            del ea
            gc.collect()
            eb = document.styles[b_name]._element
            ib = id(eb)
            if ia == ib and a_name != b_name:
                hits += 1
                if hits <= 6:
                    print("     КОЛЛИЗИЯ: %r освобождён -> %r занял ТОТ ЖЕ адрес %s" % (
                        a_name, b_name, ia))
            del eb
            gc.collect()
        except Exception:
            continue
print("  прямых коллизий адреса найдено:", hits,
      "<<< каждая = потенциальное ЛОЖНОЕ ПОПАДАНИЕ в кеш" if hits else "")

print("\n=== ТЕСТ 3: ИНСТРУМЕНТИРОВАННЫЙ _CapsResolver на полном extract() ===")
print("     (ловим случай: id уже встречался, но принадлежал ДРУГОМУ стилю)")
_orig = EX._CapsResolver._style_chain
owner_of_id = {}
false_hits = []
cycle_hits = []


def traced(self, style, attr):
    if style is not None:
        try:
            nm = style.name
        except Exception:
            nm = "?"
        key = (id(style._element), attr)
        prev = owner_of_id.get(key)
        cached = self._chain_cache.get(key, EX._MISSING)
        if prev is not None and prev != nm and cached is not EX._MISSING:
            false_hits.append((key[0], attr, prev, nm))
        owner_of_id[key] = nm
    return _orig(self, style, attr)


EX._CapsResolver._style_chain = traced
doc = EX.extract(DOC)
EX._CapsResolver._style_chain = _orig

print("  вызовов _style_chain отслежено:", len(owner_of_id))
print("  ЛОЖНЫХ ПОПАДАНИЙ (id тот же, стиль другой):", len(false_hits))
for fh in false_hits[:8]:
    print("     id=%s attr=%s : записано для %r, ЗАПРОШЕНО для %r" % fh)
if not false_hits:
    print("     (на ЭТОМ документе не выпало — мало стилей/мало churn;")
    print("      на реальном .docx из Word стилей десятки, вероятность резко выше)")

print("\n=== ТЕСТ 4: стабилен ли detection_text между повторными extract()? ===")
import hashlib
sigs = []
for attempt in range(4):
    d = EX.extract(DOC)
    rows = []
    for s in d.segments:
        dt = s.metadata.get("detection_text")
        rows.append((s.id, hashlib.sha256((dt or "").encode()).hexdigest()[:8]))
    sigs.append(hashlib.sha256(repr(rows).encode()).hexdigest()[:12])
    gc.collect()
print("  detection_text sha по 4 попыткам:", sigs)
uniq = len(set(sigs))
print("  уникальных:", uniq, "OK" if uniq == 1 else "<<< detection_text НЕСТАБИЛЕН")
n_dt = sum(1 for s in doc.segments if s.metadata.get("detection_text"))
print("  сегментов С detection_text: %d из %d" % (n_dt, len(doc.segments)))
if n_dt == 0:
    print("     ВНИМАНИЕ: в этом документе detection_text не ставится НИ РАЗУ ->")
    print("     путь _CapsResolver здесь холостой, эффект проявиться не может.")
    print("     Нужен .docx с all_caps/small_caps форматированием (реальный Word-документ).")

print("\n=== ТЕСТ 5: КОНТРОЛЬ — СТАРАЯ (id-ключ) реализация на ЭТОЙ ЖЕ фикстуре ===")
print("     (доказывает, что фикстура ЛОВИТ дефект, а не просто «зелёная»)")
_owner_old = {}
_false_old = []


def old_style_chain(self, style, attr):
    """ТОЧНАЯ копия реализации ДО фикса: ключ = id(), ссылка не удерживается."""
    if style is None:
        return None
    key = (id(style._element), attr)
    try:
        nm = style.name
    except Exception:
        nm = "?"
    prev = _owner_old.get(key)
    cached = self._chain_cache.get(key, EX._MISSING)
    if prev is not None and prev != nm and cached is not EX._MISSING:
        _false_old.append((key[0], attr, prev, nm))
    _owner_old[key] = nm
    if cached is not EX._MISSING:
        return cached
    result = None
    seen_ids = set()
    cur = style
    while cur is not None:
        el_id = id(cur._element)
        if el_id in seen_ids:
            break
        seen_ids.add(el_id)
        value = getattr(cur.font, attr)
        if value is not None:
            result = value
            break
        cur = cur.base_style
    self._chain_cache[key] = result
    return result


EX._CapsResolver._style_chain = old_style_chain
old_sigs = []
for attempt in range(4):
    d = EX.extract(DOC)
    rows = [(s.id, hashlib.sha256((s.metadata.get("detection_text") or "").encode()).hexdigest()[:8])
            for s in d.segments]
    old_sigs.append(hashlib.sha256(repr(rows).encode()).hexdigest()[:12])
    gc.collect()
EX._CapsResolver._style_chain = _orig

print("  СТАРАЯ реализация: ложных попаданий:", len(_false_old))
for fh in _false_old[:5]:
    print("     id=%s attr=%s : записано для %r, ЗАПРОШЕНО для %r" % fh)
print("  СТАРАЯ реализация: detection_text sha:", old_sigs)
print("  уникальных:", len(set(old_sigs)),
      "<<< ФИКСТУРА ЛОВИТ ДЕФЕКТ" if (len(set(old_sigs)) > 1 or _false_old)
      else "(на этой машине старая версия не сорвалась — см. примечание)")
print("  НОВАЯ реализация (выше): ложных попаданий %d, уникальных sha %d"
      % (len(false_hits), uniq))

print("\n=== ВЫВОД ===")
print("  Кеш и защита от циклов в _CapsResolver ключуются ЧИСЛОМ id() без удержания")
print("  ссылки. Прокси python-docx/lxml временные, адрес после освобождения")
print("  переиспользуется => возможны ЛОЖНОЕ ПОПАДАНИЕ и ЛОЖНЫЙ ЦИКЛ, зависящие от")
print("  момента GC => недетерминизм в одном процессе при идентичном входе.")
print("  Затрагивает detection_text (metadata) => segment.text стабилен, а L2 и L3")
print("  плавают оба. ОДИН механизм объясняет ОБА корня.")
