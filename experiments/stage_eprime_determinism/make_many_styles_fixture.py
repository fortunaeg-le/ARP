# -*- coding: utf-8 -*-
"""
E'' — фикстура «МНОГО СТИЛЕЙ» для регрессии недетерминизма _CapsResolver.

ЗАЧЕМ ОТДЕЛЬНАЯ ФИКСТУРА. Документы корпуса имеют ВСЕГО 2 стиля и физически не
могут воспроизвести дефект id()-кеша: чтобы адрес освободившегося прокси стиля был
переиспользован ДРУГИМ стилем, нужны десятки-сотни разных стилей и активный churn
прокси. На реальном договоре владельца — 1097 стилей и 1290 ложных попаданий на
2396 вызовов; на корпусе — 0 вызовов вообще. Именно бедность синтетики скрывала
дефект от всех предыдущих этапов, поэтому фикстура с большим числом стилей нужна
как ПОСТОЯННЫЙ страж.

ЧТО ВНУТРИ (данные ВЫДУМАННЫЕ, ПДн нет):
  * ~120 пользовательских стилей абзаца, сцепленных base_style в цепочки глубины
    3-4 (значение all_caps/small_caps задано НЕ у самого стиля, а у предка — чтобы
    работал именно проход по цепочке, который и кешируется);
  * часть стилей с all_caps=True, часть с small_caps=True, часть с явным False
    (явный False на верхнем уровне обязан отменять унаследованный True);
  * абзацы ЗАГЛАВНЫМ текстом с синтетическими ФИО/организациями — чтобы
    нормализация регистра реально меняла detection_text, а значит и результат
    детекции PER/ORG/ADDRESS.
"""
import os

import docx
from docx.enum.style import WD_STYLE_TYPE

ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
OUT = os.path.join(ROOT, "tests", "fixtures", "synthetic_many_styles.docx")

N_STYLES = 120

SURNAMES = ["ВОРОНЦОВ", "МАТВЕЕНКО", "ГРИНЕВИЧ", "СОЛОМИН", "ЧЕРКАСОВ",
            "ДУБРОВИН", "ЯСЕНЕВ", "КРАСНОВ", "ЖУРАВЛЁВ", "ТИХОМИРОВ"]
NAMES = ["АЛЕКСАНДР", "ДМИТРИЙ", "СЕРГЕЙ", "АНДРЕЙ", "МАКСИМ"]
PATRS = ["АЛЕКСАНДРОВИЧ", "ДМИТРИЕВИЧ", "СЕРГЕЕВИЧ", "АНДРЕЕВИЧ", "ИГОРЕВИЧ"]
ORGS = ["МЕРИДИАН", "КВАРЦИТ", "СПЕКТРУМ", "ВЕКТОР", "ГЕЛИОС"]
CITIES = ["НОВОУРАЛЬСК", "ЗАПАДНОДВИНСК", "КРАСНОКАМСК"]
STREETS = ["УЛ. ПОЛЕВАЯ", "ПР-Т МИРНЫЙ", "УЛ. ЗАРЕЧНАЯ"]


def build():
    document = docx.Document()
    styles = document.styles

    created = []
    for i in range(N_STYLES):
        name = "E2Style%03d" % i
        st = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        # цепочка: каждый 4-й стиль — корень, остальные наследуют предыдущему
        if i % 4 != 0 and created:
            try:
                st.base_style = styles[created[-1]]
            except Exception:
                pass
        # значение капса задаём ТОЛЬКО у корней цепочек — потомки наследуют,
        # значит _style_chain обязан пройти вверх (именно этот проход кешируется)
        if i % 4 == 0:
            if i % 12 == 0:
                st.font.all_caps = True
            elif i % 12 == 4:
                st.font.small_caps = True
            else:
                st.font.all_caps = False
        created.append(name)

    document.add_paragraph("ДОГОВОР № МНОГО-СТИЛЕЙ/СИНТ")
    for i in range(N_STYLES * 6):
        sname = created[i % N_STYLES]
        s = SURNAMES[i % len(SURNAMES)]
        n = NAMES[i % len(NAMES)]
        p = PATRS[i % len(PATRS)]
        o = ORGS[i % len(ORGS)]
        c = CITIES[i % len(CITIES)]
        strt = STREETS[i % len(STREETS)]
        variants = [
            "ГРАЖДАНИН %s %s %s, ИМЕНУЕМЫЙ В ДАЛЬНЕЙШЕМ СТОРОНА %d" % (s, n, p, i),
            "ООО «%s», АДРЕС: Г. %s, %s, Д. %d" % (o, c, strt, (i % 90) + 1),
            "В ЛИЦЕ ДИРЕКТОРА %s %s.%s., ДЕЙСТВУЮЩЕГО НА ОСНОВАНИИ УСТАВА" % (s, n[0], p[0]),
            "ПОДПИСЬ: _________ / %s %s %s /" % (s, n, p),
        ]
        para = document.add_paragraph(variants[i % len(variants)])
        try:
            para.style = styles[sname]
        except Exception:
            pass

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    document.save(OUT)
    print("saved:", OUT)
    print("стилей создано:", N_STYLES, "| абзацев:", N_STYLES * 6 + 1)


if __name__ == "__main__":
    build()
