# -*- coding: utf-8 -*-
"""ЧАСТЬ 3 этапа CLIENT: порождение БОЛЬШОГО документа (×10 к реальному).

Реальный документ не используется и не нужен: берём генератор синтетического
двойника этапа E'' (`experiments/stage_eprime_determinism/make_synthetic_large.py`,
все данные выдуманы, КС реквизитов валидны) и МАСШТАБИРУЕМ его.

Точка отсчёта — АУДИТ-1: `synthetic_corporate_large.txt` = 527 КБ текста,
это ~×4 к реальному договору, прогон 35 с. Значит реальный ≈ 132 КБ, а ×10 ≈
1.32 МБ. Множитель к существующей фикстуре — 2.5.

Пишет .docx И .txt-двойник: .docx — путь клиента, .txt — контроль того, что
время уходит в детекцию, а не в разбор OOXML.
"""
import os
import sys

_HERE = os.path.dirname(os.path.abspath(__file__))
_ROOT = os.path.dirname(os.path.dirname(_HERE))
sys.path.insert(1, os.path.join(_ROOT, "experiments", "stage_eprime_determinism"))

import docx                                    # noqa: E402
from docx.shared import Pt                     # noqa: E402

import make_synthetic_large as G               # noqa: E402

OUT_DIR = os.path.join(_HERE, "probe_out")
OUT_DOCX = os.path.join(OUT_DIR, "big_x10.docx")
OUT_TXT = os.path.join(OUT_DIR, "big_x10.txt")

# 2.5 × объёмы фикстуры E'' (527 КБ ≈ ×4 к реальному) => ≈ ×10 к реальному
SCALE = 2.5


def _n(base: int) -> int:
    return int(base * SCALE)


def build() -> None:
    document = docx.Document()
    document.styles["Normal"].font.size = Pt(11)

    document.add_paragraph("ДОГОВОР № КРУПНЫЙ-2026/СИНТ-X10").alignment = 1

    for i in range(_n(60)):
        p = G.full_name(male=(i % 2 == 0))
        document.add_paragraph(
            "%s, именуемый в дальнейшем «%s», действующий от имени %s," % (
                p, G.rnd.choice(["Сторона %d" % i, "Представитель"]),
                G.oblique_name_dative(i % 2 == 0)))

    for i in range(_n(1800)):
        document.add_paragraph(G.definitions_paragraph(i))
    for i in range(_n(600)):
        document.add_paragraph(G.law_paragraph(i))
    for i in range(_n(600)):
        document.add_paragraph(G.role_paragraph(i))

    G.requisite_table(document, _n(70))

    for i in range(_n(300)):
        document.add_paragraph(G.signature_block(i % 2 == 0))
        document.add_paragraph("СНИЛС %s подписанта учтён в кадровом деле." % G.gen_snils())

    for i in range(_n(1800), _n(2000)):
        document.add_paragraph(G.definitions_paragraph(i))

    os.makedirs(OUT_DIR, exist_ok=True)
    document.save(OUT_DOCX)

    with open(OUT_TXT, "w", encoding="utf-8") as f:
        f.write("\n".join(p.text for p in document.paragraphs) + "\n")

    n_par = len(document.paragraphs)
    n_cells = sum(len(r.cells) for t in document.tables for r in t.rows)
    print("docx : %s (%.1f КБ)" % (OUT_DOCX, os.path.getsize(OUT_DOCX) / 1024))
    print("txt  : %s (%.1f КБ)" % (OUT_TXT, os.path.getsize(OUT_TXT) / 1024))
    print("абзацев %d, ячеек таблиц %d, сегментов ≈ %d" % (n_par, n_cells, n_par + n_cells))


if __name__ == "__main__":
    build()
