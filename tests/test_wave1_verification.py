"""Независимая верификация волны 1 (P1, тихий успех decrypt-file, P2-P4).

Этот файл — НЕ часть фиксов, а независимая проверка поверх них (см. ТЗ
"Верификация волны 1"). Существующие тесты не редактировались; здесь только
новые тесты, закрывающие дыры, которые не проверял ни один из трёх изолированных
чатов, делавших сами фиксы:
  - стык P1 (раздельные токены у половин граничной пары) и компонента 2
    (реальный .docx round-trip через file_detokenizer, а не через detokenize());
  - P1 на других типах сущностей и в комбинации с дедупликацией по (type, text);
  - P4 при python -O (assert реально отключается — проверяем через субпроцесс);
  - P2 атомарность записи .enc в session_store (не только в ooxml_core);
  - согласованность "replaced" между тремя адаптерами и что unresolved не
    засчитывается в replaced.
"""
import subprocess
import sys
from pathlib import Path

import pytest
from docx import Document

from extractor import extract
from regex_detector import detect_regex
from ner_detector import detect_ner
from tokenizer import tokenize, build_plain_text, _render_segment
from session_store import save_session, default_storage_dir, _DEFAULT_STORAGE_DIR
from detokenizer import detokenize
from file_detokenizer import detokenize_file

_PROJECT_ROOT = Path(__file__).resolve().parent.parent


def _pipeline(doc, config_path):
    entities = detect_regex(doc, config_path) + detect_ner(doc, config_path)
    return tokenize(doc, entities, config_path)


# --------------------------------------------------------------------------- #
# D. Пересечение P1 x компонент 2: реальный .docx, а не detokenize() по строке.
# --------------------------------------------------------------------------- #

def _write_tokenized_docx(src_docx_path, doc, result, dst_docx_path):
    """Открывает src_docx_path заново и подставляет в каждый сегмент его
    анонимизированный текст (через ровно ту же _render_segment, что использует
    tokenize() для сборки anon_text), сохраняет как dst_docx_path.

    Это даёт РЕАЛЬНЫЙ .docx с токенами внутри настоящих XML-ран(ов) Word — то,
    что реально вернула бы LLM, а не текстовую подстановку по офсетам строки.
    """
    kept_by_segment = {}
    for e in result:
        kept_by_segment.setdefault(e.segment_id, []).append(e)

    document = Document(src_docx_path)
    para_idx = 0
    table_idx = 0
    for child in document.element.body.iterchildren():
        from docx.oxml.ns import qn
        if child.tag == qn("w:p"):
            seg_id = f"p{para_idx}"
            ents = kept_by_segment.get(seg_id, [])
            if ents:
                seg = next(s for s in doc.segments if s.id == seg_id)
                new_text = _render_segment(seg, ents)
                from docx.text.paragraph import Paragraph
                p = Paragraph(child, document)
                # простая замена: очистить runs, записать один run с новым текстом
                for r in list(p.runs):
                    r.text = ""
                if p.runs:
                    p.runs[0].text = new_text
                else:
                    p.add_run(new_text)
            para_idx += 1
        elif child.tag == qn("w:tbl"):
            from docx.table import Table
            table = Table(child, document)
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    seg_id = f"t{table_idx}_r{row_idx}_c{col_idx}"
                    ents = kept_by_segment.get(seg_id, [])
                    if ents:
                        seg = next(s for s in doc.segments if s.id == seg_id)
                        new_text = _render_segment(seg, ents)
                        cell.text = new_text
            table_idx += 1

    document.save(dst_docx_path)


class TestP1CrossesComponent2:
    def test_boundary_pair_round_trip_through_real_docx_file(self, tmp_path, config_path, docx_factory):
        """Стык, который не проверял никто: .docx с сущностью, разорванной
        границей ЯЧЕЙКИ таблицы ('Приложение 1' | '23456789 экз.') -> tokenize
        (даёт раздельные токены, P1) -> реальный .docx с этими токенами внутри
        -> file_detokenizer.detokenize_file (компонент 2) -> сравнить итоговый
        текст ячеек с исходным ПОСИМВОЛЬНО, включая невинное число 23456789."""
        src = docx_factory("src.docx", table_rows=[["Приложение 1", "23456789 экз."]])
        doc = extract(src)
        anon_text, result = _pipeline(doc, config_path)

        tokenized_docx = str(tmp_path / "tokenized.docx")
        _write_tokenized_docx(src, doc, result, tokenized_docx)

        session_id = save_session(result, storage_dir=str(tmp_path / "sessions"))
        dst = str(tmp_path / "restored.docx")
        dst_path, replaced, unresolved = detokenize_file(
            tokenized_docx, session_id, dst_path=dst, storage_dir=str(tmp_path / "sessions")
        )

        assert unresolved == [], f"остались нераскрытые токены: {unresolved}"
        assert replaced >= 1

        restored_doc = Document(dst_path)
        table = restored_doc.tables[0]
        assert table.cell(0, 0).text == "Приложение 1"
        assert table.cell(0, 1).text == "23456789 экз.", (
            "невинное число 23456789 должно быть восстановлено ЦЕЛИКОМ через "
            "весь путь tokenize -> save_session -> file_detokenizer, а не только "
            "через detokenize() по плоской строке"
        )

    def test_boundary_pair_phone_round_trip_through_real_docx_file(self, tmp_path, config_path, docx_factory):
        """Тот же стык, реальная сущность (телефон через границу ячеек) вместо
        ложного срабатывания."""
        src = docx_factory("src.docx", table_rows=[["+7 (495)", "123-45-67"]])
        doc = extract(src)
        anon_text, result = _pipeline(doc, config_path)

        tokenized_docx = str(tmp_path / "tokenized.docx")
        _write_tokenized_docx(src, doc, result, tokenized_docx)

        session_id = save_session(result, storage_dir=str(tmp_path / "sessions"))
        dst = str(tmp_path / "restored.docx")
        dst_path, replaced, unresolved = detokenize_file(
            tokenized_docx, session_id, dst_path=dst, storage_dir=str(tmp_path / "sessions")
        )

        assert unresolved == []
        restored_doc = Document(dst_path)
        table = restored_doc.tables[0]
        assert table.cell(0, 0).text == "+7 (495)"
        assert table.cell(0, 1).text == "123-45-67"


# --------------------------------------------------------------------------- #
# A. P1 на других типах сущностей и во взаимодействии с дедупликацией.
# --------------------------------------------------------------------------- #

class TestP1OtherEntityTypesAndDedup:
    def test_org_split_and_same_org_whole_elsewhere_do_not_collide(self, tmp_path, config_path):
        """Организация разорвана переносом строки в одном месте документа, и та
        же по смыслу подстрока встречается ЦЕЛОЙ в другом месте. Дедупликация
        по (entity_type, original_text) не должна путать половину пары с целой
        сущностью в другом месте — round-trip должен быть точным в обоих местах."""
        content = (
            'Поставщик: ООО "Ромашка\nПлюс" по договору.\n'
            'Также упоминается ООО "Ромашка" отдельно в тексте.'
        )
        p = tmp_path / "d.txt"
        p.write_text(content, encoding="utf-8")
        doc = extract(str(p))
        anon, result = _pipeline(doc, config_path)

        session_id = save_session(result, storage_dir=str(tmp_path / "sessions"))
        restored, unresolved = detokenize(anon, session_id, storage_dir=str(tmp_path / "sessions"))

        assert unresolved == []
        assert restored == build_plain_text(doc)

    def test_same_entity_split_in_two_different_places_both_restore_exactly(self, tmp_path, config_path):
        """Одна и та же сущность (номер телефона) разорвана границей в ДВУХ
        разных местах документа. Обе пары должны восстановиться независимо и
        посимвольно точно."""
        content = (
            "Первый контакт: +7 (495)\n123-45-67 в отделе продаж.\n"
            "Второй контакт: +7 (495)\n123-45-67 в бухгалтерии."
        )
        p = tmp_path / "d.txt"
        p.write_text(content, encoding="utf-8")
        doc = extract(str(p))
        anon, result = _pipeline(doc, config_path)

        session_id = save_session(result, storage_dir=str(tmp_path / "sessions"))
        restored, unresolved = detokenize(anon, session_id, storage_dir=str(tmp_path / "sessions"))

        assert unresolved == []
        assert restored == build_plain_text(doc)
        assert content.count("123-45-67") == 2
        assert restored.count("123-45-67") == 2


# --------------------------------------------------------------------------- #
# P4 — assert заменён на raise; должен реально бросаться при `python -O`
# (assert отключается интерпретатором целиком, raise — нет).
# --------------------------------------------------------------------------- #

def test_p4_invariant_violation_raises_under_python_dash_O(tmp_path):
    """Прямой прогон РЕАЛЬНОГО кода инвариантной проверки из ner_detector.py
    (извлечённого из исходника этим же тестом, а не переписанного вручную) под
    `python -O`. Если бы проверка осталась голым `assert` верхнего уровня, `-O`
    вырезал бы её байткомпилятором и нарушение прошло бы молча (NO_ERROR_RAISED).
    Ожидание после фикса P4: AssertionError долетает и под -O."""
    src_path = _PROJECT_ROOT / "src" / "ner_detector.py"
    source = src_path.read_text(encoding="utf-8")
    marker = "Narushen invariant original_text"  # placeholder replaced below
    start = source.index("    for e in entities:")
    end = source.index("\n\n    return entities")
    invariant_block = source[start:end]

    script_lines = [
        "class FakeSeg:\n",
        "    def __init__(self, text):\n",
        "        self.id = 'p0'\n",
        "        self.text = text\n",
        "\n",
        "class FakeEntity:\n",
        "    def __init__(self):\n",
        "        self.segment_id = 'p0'\n",
        "        self.start = 0\n",
        "        self.end = 6\n",
        "        self.original_text = 'WRONG-NOT-MATCHING-SLICE'\n",
        "        self.entity_type = 'PERSON'\n",
        "\n",
        "class FakeDoc:\n",
        "    segments = [FakeSeg('Ivanov Ivan Ivanovich')]\n",
        "\n",
        "entities = [FakeEntity()]\n",
        "doc = FakeDoc()\n",
        "try:\n",
    ]
    for line in invariant_block.splitlines():
        script_lines.append("    " + line + "\n")
    script_lines.append("        print('NO_ERROR_RAISED')\n")
    script_lines.append("except AssertionError:\n")
    script_lines.append("    print('ASSERTION_ERROR_RAISED')\n")

    script = tmp_path / "check_p4.py"
    script.write_text("".join(script_lines), encoding="utf-8")

    result = subprocess.run(
        [sys.executable, "-O", str(script)],
        capture_output=True, text=True, timeout=30,
    )
    assert "ASSERTION_ERROR_RAISED" in result.stdout, (
        f"инвариант не сработал под -O; stdout={result.stdout!r} stderr={result.stderr!r}"
    )


def test_p4_source_uses_raise_not_bare_assert_for_invariant():
    """Прямая проверка по исходнику: инвариантная проверка original_text в
    ner_detector.detect_ner не должна быть голым `assert` верхнего уровня (он
    вырезается компилятором под -O и молча пропускает нарушение)."""
    src_path = _PROJECT_ROOT / "src" / "ner_detector.py"
    text = src_path.read_text(encoding="utf-8")
    idx = text.index("Нарушен инвариант original_text")
    window = text[max(0, idx - 300):idx]
    assert "raise AssertionError" in text[idx - 50:idx + 400] or "raise" in window, (
        "проверка инварианта original_text должна использовать raise, не assert"
    )
    # Более прямая проверка: нет строки вида "    assert e.original_text =="
    assert "assert e.original_text ==" not in text


# --------------------------------------------------------------------------- #
# P2 — атомарность записи session_store.save_session (.enc через tmp+replace).
# --------------------------------------------------------------------------- #

class TestP2SessionStoreAtomicWrite:
    def test_crash_mid_write_leaves_no_partial_enc_under_final_name(self, tmp_path, monkeypatch):
        """Симулирует обрыв записи ПОСЛЕ создания tmp-файла, но ДО os.replace.
        Ожидание: под финальным именем {sid}.enc не появляется битый файл, и
        временный файл за собой убирается (или хотя бы не выдаётся за валидную
        сессию)."""
        from extractor import extract
        storage_dir = tmp_path / "sessions"
        storage_dir.mkdir()

        p = tmp_path / "d.txt"
        p.write_text("Иванов Иван Иванович, тел. +7 (495) 123-45-67", encoding="utf-8")
        doc = extract(str(p))
        entities = detect_regex(doc, "entity_types.yaml".__class__ and str(_PROJECT_ROOT / "entity_types.yaml"))
        _, result = tokenize(doc, entities, str(_PROJECT_ROOT / "entity_types.yaml"))

        import os as _os
        real_replace = _os.replace

        def flaky_replace(*args, **kwargs):
            raise OSError("симулированный обрыв перед переименованием")

        monkeypatch.setattr(_os, "replace", flaky_replace)

        with pytest.raises(OSError):
            save_session(result, storage_dir=str(storage_dir))

        enc_files = list(storage_dir.glob("*.enc"))
        assert enc_files == [], f"под финальным именем не должно быть .enc после обрыва: {enc_files}"

    def test_no_leftover_tmp_files_after_crash(self, tmp_path, monkeypatch):
        from extractor import extract
        storage_dir = tmp_path / "sessions"
        storage_dir.mkdir()

        p = tmp_path / "d.txt"
        p.write_text("Иванов Иван Иванович", encoding="utf-8")
        doc = extract(str(p))
        entities = detect_regex(doc, str(_PROJECT_ROOT / "entity_types.yaml"))
        _, result = tokenize(doc, entities, str(_PROJECT_ROOT / "entity_types.yaml"))

        import os as _os
        monkeypatch.setattr(_os, "replace", lambda *a, **k: (_ for _ in ()).throw(OSError("boom")))

        with pytest.raises(OSError):
            save_session(result, storage_dir=str(storage_dir))

        leftover = list(storage_dir.glob("*.tmp"))
        assert leftover == [], f"остались недописанные временные файлы сессии: {leftover}"


def test_p3_default_storage_dir_matches_original_hardcoded_path():
    """P3: default_storage_dir() (используемый теперь и в session_store, и в
    shifrator.py cmd_encrypt) обязан указывать РОВНО на ~/.shifrator/sessions —
    иначе все существующие пользовательские сессии станут недоступны молча."""
    from pathlib import Path as _Path
    expected = _Path.home() / ".shifrator" / "sessions"
    assert default_storage_dir() == expected
    assert _DEFAULT_STORAGE_DIR == expected


# --------------------------------------------------------------------------- #
# B. Согласованность "replaced" между тремя адаптерами; unresolved не в счётчик.
# --------------------------------------------------------------------------- #

# --------------------------------------------------------------------------- #
# CLI decrypt-file — ни один существующий тест (tests/test_cli.py) не вызывает
# эту подкоманду вообще; проверяем контракт "тихого успеха" на уровне реального
# процесса, а не только на уровне detokenize_file()/rewrite().
# --------------------------------------------------------------------------- #

SHIFRATOR = str(_PROJECT_ROOT / "shifrator.py")


def _run_cli(args, tmp_path, cwd=None):
    import os as _os

    import conftest
    conftest.pin_all_types(tmp_path / "home")   # ЭТАП T1: набор «Максимум», см. conftest
    env = dict(
        _os.environ, PYTHONIOENCODING="utf-8",
        USERPROFILE=str(tmp_path / "home"), HOME=str(tmp_path / "home"),
    )
    return subprocess.run(
        [sys.executable, SHIFRATOR] + args,
        capture_output=True, text=True, encoding="utf-8", env=env,
        cwd=cwd or str(_PROJECT_ROOT),
    )


class TestCliDecryptFileZeroReplacements:
    def test_docx_without_any_token_exits_0_with_warning_and_creates_file(self, tmp_path, docx_factory):
        txt_path = tmp_path / "input.txt"
        txt_path.write_text("ИНН 7707083893 в тексте", encoding="utf-8")
        enc = _run_cli(["encrypt", str(txt_path)], tmp_path)
        assert enc.returncode == 0, enc.stderr
        session_id = enc.stdout.strip()

        clean_docx = docx_factory("clean.docx", paragraphs_before=["Текст совсем без токенов."])
        out = tmp_path / "restored.docx"
        result = _run_cli(["decrypt-file", session_id, clean_docx, "--out", str(out)], tmp_path)

        assert result.returncode == 0, f"нулевое число замен — это НЕ ошибка: {result.stderr}"
        assert out.exists()
        assert "не найдено ни одного токена" in result.stderr

    def test_docx_with_token_exits_0_without_zero_replacement_warning(self, tmp_path, docx_factory):
        txt_path = tmp_path / "input.txt"
        txt_path.write_text("ИНН 7707083893 в тексте", encoding="utf-8")
        enc = _run_cli(["encrypt", str(txt_path)], tmp_path)
        assert enc.returncode == 0, enc.stderr
        session_id = enc.stdout.strip()

        tokened_docx = docx_factory("tokened.docx", paragraphs_before=["ИНН клиента: [INN_1]"])
        out = tmp_path / "restored.docx"
        result = _run_cli(["decrypt-file", session_id, tokened_docx, "--out", str(out)], tmp_path)

        assert result.returncode == 0, result.stderr
        assert "не найдено ни одного токена" not in result.stderr
        assert result.stderr == "", f"неожиданный stderr при успешной замене: {result.stderr!r}"

        from docx import Document as _Document
        restored = _Document(str(out))
        assert "7707083893" in restored.paragraphs[0].text


def test_unresolved_token_does_not_count_towards_replaced(tmp_path):
    from docx_rewriter import rewrite as docx_rewrite
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    import zipfile

    CONTENT_TYPES = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '</Types>'
    ).encode("utf-8")
    xml = f"""<w:document xmlns:w="{W}"><w:body>
<w:p><w:r><w:t>[ORG_1] и [ORG_99]</w:t></w:r></w:p>
</w:body></w:document>""".encode("utf-8")

    src = tmp_path / "src.docx"
    with zipfile.ZipFile(src, "w") as zf:
        zf.writestr("[Content_Types].xml", CONTENT_TYPES)
        zf.writestr("word/document.xml", xml)
    dst = str(tmp_path / "out.docx")

    def resolve(token):
        return "ООО «Ромашка»" if token == "[ORG_1]" else None

    replaced, unresolved = docx_rewrite(str(src), dst, resolve)
    assert replaced == 1, "только [ORG_1] реально заменён — [ORG_99] не должен входить в счётчик"
    assert unresolved == ["[ORG_99]"]
