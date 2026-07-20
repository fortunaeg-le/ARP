"""Тесты фикса утечки ПДн через нестандартный регистр (all_caps / small_caps /
строчный ввод). См. docs/SHIFRATOR_SPEC_AI.md, блок 1 (detection_text) и ТЗ
"Фикс утечки ПДн через нестандартный регистр".

Механизм: рядом с настоящим segment.text живёт detection_text — копия ТОЙ ЖЕ
ДЛИНЫ с нормализованным регистром. Детекторы ищут в detection_text, но
original_text вырезают из настоящего text по тем же оффсетам. Равная длина —
фундамент; при её нарушении detection_text не создаётся вообще.

Использует реальную natasha через session-scope фикстуру ner_detector_module.
"""
import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from extractor import extract, _normalize_case, _MAX_HEURISTIC_LEN
from regex_detector import detect_regex
from tokenizer import tokenize, build_plain_text
from session_store import save_session
from detokenizer import detokenize


# --------------------------------------------------------------------------- #
# Хелперы построения .docx с форматированием ран (docx_factory этого не умеет).
# --------------------------------------------------------------------------- #

def _make_docx(path, paragraphs):
    """paragraphs: list[list[(text, fmt)]], fmt = {} | {"all_caps": True} |
    {"small_caps": True}. Каждый внутренний список — один параграф из ран."""
    doc = Document()
    for runs in paragraphs:
        p = doc.add_paragraph()
        for text, fmt in runs:
            r = p.add_run(text)
            if fmt.get("all_caps"):
                r.font.all_caps = True
            if fmt.get("small_caps"):
                r.font.small_caps = True
    doc.save(str(path))
    return str(path)


def _seg(doc, seg_id):
    return next(s for s in doc.segments if s.id == seg_id)


# --------------------------------------------------------------------------- #
# Изменение 1 — detection_text из форматирования all_caps / small_caps.
# --------------------------------------------------------------------------- #

class TestDocxAllCaps:
    def test_all_caps_person_found_original_stays_lowercase(self, tmp_path, config_path, ner_detector_module):
        """ФИО строчными в XML, визуально заглавные через all_caps -> PERSON найден,
        а original_text сохранён СТРОЧНЫМ (как в оригинале)."""
        src = _make_docx(tmp_path / "d.docx", [
            [("стороны: ", {}), ("иванов иван иванович", {"all_caps": True})],
        ])
        doc = extract(src)
        seg = _seg(doc, "p0")
        assert seg.metadata["detection_text"] == "стороны: Иванов Иван Иванович"

        persons = [e for e in ner_detector_module.detect_ner(doc, config_path)
                   if e.entity_type == "PERSON"]
        assert persons, "ФИО под all_caps должно находиться после нормализации"
        assert any(p.original_text == "иванов иван иванович" for p in persons), (
            f"original_text обязан быть строчным оригиналом, получено: "
            f"{[p.original_text for p in persons]}"
        )

    def test_small_caps_person_found(self, tmp_path, config_path, ner_detector_module):
        """small_caps -> аналогично all_caps."""
        src = _make_docx(tmp_path / "d.docx", [
            [("стороны: ", {}), ("петров пётр петрович", {"small_caps": True})],
        ])
        doc = extract(src)
        assert _seg(doc, "p0").metadata["detection_text"] == "стороны: Петров Пётр Петрович"
        persons = [e for e in ner_detector_module.detect_ner(doc, config_path)
                   if e.entity_type == "PERSON"]
        assert any(p.original_text == "петров пётр петрович" for p in persons)

    def test_legal_form_abbreviation_stays_uppercase(self, tmp_path, config_path, ner_detector_module):
        """'ИП ПИРОГОВА А.С.' через all_caps (в XML строчными) -> после нормализации
        'ИП' остаётся 'ИП', а не 'Ип'; сущности найдены."""
        src = _make_docx(tmp_path / "d.docx", [
            [("ип пирогова а.с.", {"all_caps": True})],
        ])
        doc = extract(src)
        assert _seg(doc, "p0").metadata["detection_text"] == "ИП Пирогова А.С."

        persons = [e for e in ner_detector_module.detect_ner(doc, config_path)
                   if e.entity_type == "PERSON"]
        assert persons, "ФИО при юр-форме ИП должно находиться"
        # original_text любой сущности — строго срез строчного оригинала.
        for p in persons:
            assert p.original_text == "ип пирогова а.с."[p.start:p.end]


class TestDocxNoNormalizationNeeded:
    def test_long_lowercase_unformatted_paragraph_gets_no_detection_text(self, tmp_path):
        """Длинный сегмент строчными БЕЗ форматирования и БЕЗ нумерации ->
        detection_text НЕ создаётся (известное ограничение, см. раздел про дыру)."""
        long_lower = ("в лице генерального директора иванова ивана ивановича, "
                      "действующего на основании устава, именуемое в дальнейшем "
                      "сторона по настоящему договору поставки продукции, а также "
                      "по всем сопутствующим соглашениям и приложениям к нему")
        assert len(long_lower) > _MAX_HEURISTIC_LEN
        src = _make_docx(tmp_path / "d.docx", [[(long_lower, {})]])
        doc = extract(src)
        assert "detection_text" not in _seg(doc, "p0").metadata

    def test_normal_mixed_case_segment_has_no_detection_text(self, tmp_path, config_path, ner_detector_module):
        """Обычный сегмент без проблем -> detection_text отсутствует, поведение
        идентично текущему (регрессия)."""
        src = _make_docx(tmp_path / "d.docx", [
            [("Договор подписал Иванов И.И. от лица ООО «Ромашка».", {})],
        ])
        doc = extract(src)
        assert "detection_text" not in _seg(doc, "p0").metadata
        # Поведение детекции — как раньше: PERSON находит detect_ner. ORG сменил
        # источник (этап A): его даёт структурный движок anchor_registry.detect_org,
        # а не Natasha-ORG. Проверяем оба: PERSON из NER, ORG из движка.
        from anchor_registry import detect_org
        types = {e.entity_type for e in ner_detector_module.detect_ner(doc, config_path)}
        assert "PERSON" in types
        assert any(e.entity_type == "ORG" for e in detect_org(doc))


# --------------------------------------------------------------------------- #
# Изменение 2 — узкая эвристика для настоящего строчного ввода (пункт перечисления).
# --------------------------------------------------------------------------- #

class TestLowercaseHeuristic:
    def test_numbered_lowercase_item_gets_detection_text_and_person(self, txt_factory, config_path, ner_detector_module):
        """Короткий '(1) иванов иван иванович' без форматирования (.txt) ->
        Изменение 2 срабатывает, Entity найден."""
        src = txt_factory("(1) иванов иван иванович")
        doc = extract(src)
        assert _seg(doc, "l0").metadata["detection_text"] == "(1) Иванов Иван Иванович"

        persons = [e for e in ner_detector_module.detect_ner(doc, config_path)
                   if e.entity_type == "PERSON"]
        assert any(p.original_text == "иванов иван иванович" for p in persons)

    def test_numbered_dot_form_also_matches(self, txt_factory):
        """Форма '1.' тоже распознаётся как пункт перечисления."""
        src = txt_factory("1. сидоров сидор сидорович")
        doc = extract(src)
        assert _seg(doc, "l0").metadata["detection_text"] == "1. Сидоров Сидор Сидорович"

    def test_unnumbered_lowercase_line_now_gets_detection_text(self, txt_factory):
        """ЭТАП 2b: строчная строка БЕЗ нумерации теперь нормализуется.

        До 2b эвристика требовала признак пункта перечисления ("1."/"(1)"), и это
        была ЗАДОКУМЕНТИРОВАННАЯ непокрытая дыра: ФИО строчными в обычном абзаце
        не находилось (recall на сплошь строчных ФИО был РОВНО 0.0%). Требование
        нумерации снято — тест перевёрнут вместе с поведением."""
        src = txt_factory("иванов иван иванович подписал договор")
        doc = extract(src)
        assert (_seg(doc, "l0").metadata["detection_text"]
                == "Иванов Иван Иванович Подписал Договор")


# --------------------------------------------------------------------------- #
# Этап 2b — регистровая нормализация на уровне СЛОВА и её щиты.
# --------------------------------------------------------------------------- #

class TestStage2bWordLevelCase:
    def test_uppercase_name_inside_mixed_segment(self, txt_factory, config_path,
                                                 ner_detector_module):
        """ЗАГЛАВНОЕ ФИО в сегменте СМЕШАННОГО регистра. Посегментная эвристика
        мимо этого проходила (isupper() у сегмента ложно) — словная ловит."""
        src = txt_factory("Исполнитель: ПЕТРОВА ОЛЬГА ИВАНОВНА")
        doc = extract(src)
        assert (_seg(doc, "l0").metadata["detection_text"]
                == "Исполнитель: Петрова Ольга Ивановна")
        persons = [e for e in ner_detector_module.detect_ner(doc, config_path)
                   if e.entity_type == "PERSON"]
        assert any("ПЕТРОВА" in p.original_text for p in persons)

    def test_alternating_case_word_is_normalized(self, txt_factory):
        """'пЕтРоВ' — заглавная после строчной внутри слова, так не пишут."""
        src = txt_factory("Подписал кУзНеЦоВа мАрИя")
        doc = extract(src)
        assert (_seg(doc, "l0").metadata["detection_text"]
                == "Подписал Кузнецова Мария")

    def test_normal_case_text_is_left_byte_for_byte(self, txt_factory):
        """ГЛАВНАЯ ГАРАНТИЯ: на тексте с нормальным регистром преобразование
        ТОЖДЕСТВЕННО, detection_text вообще не создаётся. Именно это защищает
        уже работающую детекцию (PER recall на смешанном регистре) от 2b."""
        src = txt_factory("Договор подписан. Стороны: Иванов И. И. и ООО «Ромашка».")
        doc = extract(src)
        assert "detection_text" not in _seg(doc, "l0").metadata

    def test_legal_form_abbreviations_stay_uppercase(self, txt_factory):
        """Слова 2-3 букв (ООО/ИП/АО) — юр-формы, а не капс: остаются заглавными,
        иначе рушится синтаксический слой составных сущностей ('ИП + ФИО')."""
        src = txt_factory("ип пирогова а.с.")
        doc = extract(src)
        assert _seg(doc, "l0").metadata["detection_text"] == "ИП Пирогова А.С."

    def test_numeric_token_case_is_untouched(self, txt_factory):
        """ЩИТ РЕКВИЗИТА. Свод цифровых омоглифов в normalizer регистрозависим
        ('б'->6 есть, 'Б'->6 НЕТ). Поднять регистр у БИК, где буквы стоят вместо
        цифр, значит сломать свод и потерять реквизит ЦЕЛИКОМ. Замер этапа 2b
        поймал ровно это: 4 реквизита теряли маску полностью."""
        # ФИО рядом — чтобы detection_text вообще создался (без единого изменения
        # он не создаётся): так видно, что соседний реквизит при этом прошёл мимо
        # нормализации байт-в-байт. Три заглавных слова — иначе не прогон.
        src = txt_factory("ПЕТРОВА ОЛЬГА ИВАНОВНА бик 04Ч525ЗбО счёт 4О80281О4бІ704803064")
        doc = extract(src)
        det = _seg(doc, "l0").metadata["detection_text"]
        assert det.startswith("Петрова ")
        assert "04Ч525ЗбО" in det and "4О80281О4бІ704803064" in det

    def test_invisible_chars_do_not_split_word(self, txt_factory):
        """Слово читается СКВОЗЬ невидимые (Cf), как и в normalizer: иначе
        'о<SHY>лЬгА' распалось бы на 'о' и 'лЬгА', одиночная 'о' не выглядела бы
        аномальной, и имя чинилось бы наполовину ('о Льга')."""
        src = txt_factory("Получатель о\xadлЬгА")
        doc = extract(src)
        assert _seg(doc, "l0").metadata["detection_text"] == "Получатель О\xadльга"

    def test_lone_uppercase_acronym_is_not_touched(self, txt_factory):
        """ОДИНОКОЕ заглавное слово среди нормального текста — аббревиатура, а не
        ФИО. Приведённое к Титульному, 'ОКВЭД' -> 'Оквэд' читается Наташей как
        фамилия: замер дал +17 ложных PER на негативах, 13 из них ровно 'ОКВЭД'."""
        src = txt_factory("СТ 7.32-2017, артикул RM-7584-29, код ОКВЭД")
        doc = extract(src)
        assert "detection_text" not in _seg(doc, "l0").metadata

    def test_uppercase_run_of_three_is_normalized(self, txt_factory):
        """Контроль к предыдущему: ФИО заглавными — ПРОГОН из >=3 заглавных слов,
        он чинится. Инициалы считаются частью прогона ('КРЫЛОВА Е. П.')."""
        doc = extract(txt_factory("Подписал КРЫЛОВА Е. П."))
        assert _seg(doc, "l0").metadata["detection_text"] == "Подписал Крылова Е. П."

    def test_two_word_caps_run_is_left_alone(self, txt_factory):
        """Прогон из ДВУХ заглавных слов не трогаем: чаще это топоним без
        маркеров, чем ФИО. Замер: 'МОСКВА ЛЕНИНА 47/3 КВ 67' до 2b случайно
        попадал под NER-маску целиком, а нормализация регистра эту случайную
        маску убирала и ОТКРЫВАЛА адрес — 2b создавал новую утечку. Цена порога —
        двусловные заглавные ФИО не чинятся; инвариант дороже."""
        doc = extract(txt_factory("МОСКВА ЛЕНИНА 47/3 КВ 67"))
        assert "detection_text" not in _seg(doc, "l0").metadata

    def test_combo_segment_is_left_alone_entirely(self, txt_factory):
        """Stage2b-A: сегмент со словом СМЕШАННОГО алфавита (омоглифная подмена
        латиницей) этап 2b не трогает ВОВСЕ — ни одного слова.

        Щит стоит на сегменте, а не на слове, и это принципиально: словной щит
        замерен и отвергнут (полунормализация — чистые слова имени к Титульному,
        омоглифное рваным — теряла 7 ФИО из-под маски ЦЕЛИКОМ). Инвариант 2b:
        не создать НИ ОДНОЙ новой утечки; combo-стык вынесен отдельным шагом."""
        # 'иBаНоB' — латинская B внутри кириллицы; 'пЕтРоВа' рядом чистая.
        src = txt_factory("иBаНоB пЕтРоВа")
        doc = extract(src)
        assert "detection_text" not in _seg(doc, "l0").metadata

    def test_pure_cyrillic_segment_still_normalized(self, txt_factory):
        """Контроль к предыдущему: без латиницы тот же текст чинится — то есть
        щит сужает 2b ровно на combo, а не отключает его."""
        src = txt_factory("иванов пЕтРоВа")
        doc = extract(src)
        assert _seg(doc, "l0").metadata["detection_text"] == "иванов Петрова"

    def test_length_is_preserved_by_word_level_branch(self, txt_factory):
        """Инвариант равной длины — фундамент схемы оффсетов — держится и на
        словной ветви, включая невидимые и реквизиты."""
        line = "ПЕТРОВА О\xadЛЬГА, бик 04Ч525ЗбО, кУзНеЦоВа"
        doc = extract(txt_factory(line))
        seg = _seg(doc, "l0")
        assert len(seg.metadata["detection_text"]) == len(seg.text)


# --------------------------------------------------------------------------- #
# Жёсткий инвариант равной длины.
# --------------------------------------------------------------------------- #

class TestLengthInvariant:
    def test_every_detection_text_equals_source_length(self, tmp_path, txt_factory):
        """Для КАЖДОГО созданного detection_text len(detection_text) == len(text)."""
        src = _make_docx(tmp_path / "d.docx", [
            [("иванов иван иванович", {"all_caps": True})],
            [("петров пётр петрович", {"small_caps": True})],
            [("ип пирогова а.с.", {"all_caps": True})],
            [("обычный текст без флагов", {})],
        ])
        doc = extract(src)
        txt = txt_factory("(1) сидоров сидор сидорович")
        doc2 = extract(txt)

        created = 0
        for d in (doc, doc2):
            for seg in d.segments:
                if "detection_text" in seg.metadata:
                    created += 1
                    assert len(seg.metadata["detection_text"]) == len(seg.text), (
                        f"нарушен инвариант длины в {seg.id}"
                    )
        assert created >= 4, "ожидались хотя бы 4 сегмента с detection_text"

    def test_length_changing_normalization_is_rejected_and_warned(self, tmp_path, capsys):
        """Сегмент, где нормализация изменила бы длину ('ß'.upper() == 'SS'), под
        all_caps -> detection_text НЕ создан, предупреждение залогировано."""
        # Санити: одиночная 'ß' действительно меняет длину при нашей нормализации.
        assert len(_normalize_case("ß")) != len("ß")

        src = _make_docx(tmp_path / "d.docx", [[("ß", {"all_caps": True})]])
        doc = extract(src)
        assert "detection_text" not in _seg(doc, "p0").metadata
        err = capsys.readouterr().err
        assert "detection_text не создан" in err


# --------------------------------------------------------------------------- #
# Round-trip: нормализованная версия НЕ должна попасть в сессию.
# --------------------------------------------------------------------------- #

class TestRoundTrip:
    def test_all_caps_person_round_trip_char_for_char(self, tmp_path, config_path, ner_detector_module):
        """encrypt -> decrypt на .docx с all_caps ФИО -> восстановленный текст
        ПОСИМВОЛЬНО равен оригиналу (detection_text в сессию не попадает)."""
        src = _make_docx(tmp_path / "d.docx", [
            [("стороны: ", {}), ("иванов иван иванович", {"all_caps": True})],
            [("инн 7707083893", {"all_caps": True})],
        ])
        doc = extract(src)
        entities = detect_regex(doc, config_path) + ner_detector_module.detect_ner(doc, config_path)
        anon, result = tokenize(doc, entities, config_path)

        session_id = save_session(result, storage_dir=str(tmp_path / "sessions"))
        restored, unresolved = detokenize(anon, session_id, storage_dir=str(tmp_path / "sessions"))

        assert unresolved == []
        assert restored == build_plain_text(doc), (
            "восстановленный текст обязан посимвольно совпасть с исходным "
            "(строчным) — нормализованная версия НЕ должна попасть в сессию"
        )
        # Явно: анонимизированный текст содержит токены, а восстановленный — строчный оригинал.
        assert "иванов иван иванович" in restored


# --------------------------------------------------------------------------- #
# Дофикс: капс, унаследованный от СТИЛЯ (абзаца/символа/таблицы/base_style),
# а не только прямое форматирование рана. В OOXML all_caps трёхзначен
# (True/False/None-наследуй); resolver проходит цепочку наследования.
# --------------------------------------------------------------------------- #

def _add_para_style(doc, name, *, all_caps=None, small_caps=None, base=None):
    st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    if base is not None:
        st.base_style = base
    if all_caps is not None:
        st.font.all_caps = all_caps
    if small_caps is not None:
        st.font.small_caps = small_caps
    return st


def _add_char_style(doc, name, *, all_caps=None, small_caps=None):
    st = doc.styles.add_style(name, WD_STYLE_TYPE.CHARACTER)
    if all_caps is not None:
        st.font.all_caps = all_caps
    if small_caps is not None:
        st.font.small_caps = small_caps
    return st


class TestStyleInheritedCaps:
    def test_paragraph_style_caps_found(self, tmp_path, config_path, ner_detector_module):
        """ФИО строчными, капс через СТИЛЬ АБЗАЦА -> Entity найден, original строчный."""
        path = tmp_path / "d.docx"
        doc = Document()
        st = _add_para_style(doc, "CapsPara", all_caps=True)
        p = doc.add_paragraph()
        p.style = st
        p.add_run("иванов иван иванович")
        doc.save(str(path))

        d = extract(str(path))
        assert _seg(d, "p0").metadata["detection_text"] == "Иванов Иван Иванович"
        persons = [e for e in ner_detector_module.detect_ner(d, config_path)
                   if e.entity_type == "PERSON"]
        assert any(pe.original_text == "иванов иван иванович" for pe in persons)

    def test_character_style_caps_found(self, tmp_path, config_path, ner_detector_module):
        """Капс через СТИЛЬ СИМВОЛА (character style) на ране -> найден."""
        path = tmp_path / "d.docx"
        doc = Document()
        cs = _add_char_style(doc, "CapsChar", all_caps=True)
        p = doc.add_paragraph()
        r = p.add_run("петров пётр петрович")
        r.style = cs
        doc.save(str(path))

        d = extract(str(path))
        assert _seg(d, "p0").metadata["detection_text"] == "Петров Пётр Петрович"
        persons = [e for e in ner_detector_module.detect_ner(d, config_path)
                   if e.entity_type == "PERSON"]
        assert any(pe.original_text == "петров пётр петрович" for pe in persons)

    def test_base_style_chain_caps_found(self, tmp_path, config_path, ner_detector_module):
        """Капс на РОДИТЕЛЬСКОМ стиле, применён стиль-ребёнок без своего капса
        (цепочка base_style) -> найден."""
        path = tmp_path / "d.docx"
        doc = Document()
        parent = _add_para_style(doc, "ParentCaps", all_caps=True)
        child = _add_para_style(doc, "ChildPlain", base=parent)
        p = doc.add_paragraph()
        p.style = child
        p.add_run("сидоров сидор сидорович")
        doc.save(str(path))

        d = extract(str(path))
        assert _seg(d, "p0").metadata["detection_text"] == "Сидоров Сидор Сидорович"
        persons = [e for e in ner_detector_module.detect_ner(d, config_path)
                   if e.entity_type == "PERSON"]
        assert any(pe.original_text == "сидоров сидор сидорович" for pe in persons)

    def test_table_style_caps_found(self, tmp_path, config_path, ner_detector_module):
        """Капс через СТИЛЬ ТАБЛИЦЫ -> ФИО в ячейке найдено (python-docx не отражает
        это в run/paragraph style, резолвер берёт table.style напрямую)."""
        path = tmp_path / "d.docx"
        doc = Document()
        ts = doc.styles.add_style("CapsTable", WD_STYLE_TYPE.TABLE)
        ts.font.all_caps = True
        t = doc.add_table(rows=1, cols=1)
        t.style = ts
        t.cell(0, 0).paragraphs[0].add_run("кузнецов иван петрович")
        doc.save(str(path))

        d = extract(str(path))
        cell_seg = _seg(d, "t0_r0_c0")
        assert cell_seg.metadata["detection_text"] == "Кузнецов Иван Петрович"
        persons = [e for e in ner_detector_module.detect_ner(d, config_path)
                   if e.entity_type == "PERSON"]
        assert any(pe.original_text == "кузнецов иван петрович" for pe in persons)

    def test_small_caps_via_paragraph_style_found(self, tmp_path, config_path, ner_detector_module):
        """small_caps через стиль абзаца -> аналогично all_caps."""
        path = tmp_path / "d.docx"
        doc = Document()
        st = _add_para_style(doc, "SmallCapsPara", small_caps=True)
        p = doc.add_paragraph()
        p.style = st
        p.add_run("николаев николай николаевич")
        doc.save(str(path))

        d = extract(str(path))
        assert _seg(d, "p0").metadata["detection_text"] == "Николаев Николай Николаевич"


class TestStyleCapsCancellation:
    # Текст проб здесь СМЕШАННОГО регистра ('иванов Иван Иванович'), и это
    # существенно. Проверяется механизм Изменения 1 (наследование all_caps по
    # цепочке стилей), а его признак — сам факт появления detection_text. После
    # этапа 2b сплошь строчный сегмент нормализуется УЖЕ ПО ДРУГОЙ ветви
    # (регистровая эвристика), поэтому на строчной пробе «detection_text нет»
    # перестало означать «капс отменён» — проба смешала бы два механизма.
    # Смешанный регистр для ветви 2b тождественен (она правит только заглавные
    # слова и слова с внутренним флипом), значит ключ появится ТОЛЬКО если
    # сработает наследование капса — ровно то, что тест и должен ловить.
    def test_direct_run_false_cancels_inherited_true(self, tmp_path):
        """Явный False на ране поверх унаследованного True (стиль абзаца) ->
        detection_text НЕ создаётся (трёхзначная логика не схлопнута в двузначную)."""
        path = tmp_path / "d.docx"
        doc = Document()
        st = _add_para_style(doc, "CapsPara", all_caps=True)
        p = doc.add_paragraph()
        p.style = st
        r = p.add_run("иванов Иван Иванович")
        r.font.all_caps = False   # отмена унаследованного капса
        doc.save(str(path))

        d = extract(str(path))
        assert "detection_text" not in _seg(d, "p0").metadata

    def test_child_style_false_cancels_parent_true(self, tmp_path):
        """Явный False на стиле-ребёнке отменяет True базового стиля -> нет капса."""
        path = tmp_path / "d.docx"
        doc = Document()
        parent = _add_para_style(doc, "ParentCaps", all_caps=True)
        child = _add_para_style(doc, "ChildOff", all_caps=False, base=parent)
        p = doc.add_paragraph()
        p.style = child
        p.add_run("иванов Иван Иванович")
        doc.save(str(path))

        d = extract(str(path))
        assert "detection_text" not in _seg(d, "p0").metadata


class TestStyleCapsRobustness:
    def test_cyclic_base_style_does_not_hang_or_crash(self, tmp_path):
        """Битый .docx с циклической цепочкой base_style (A->B->A) -> извлечение
        завершается, не зависает и не падает; безопасная деградация."""
        path = tmp_path / "d.docx"
        doc = Document()
        a = _add_para_style(doc, "StyleA")
        b = _add_para_style(doc, "StyleB")
        a.base_style = b
        b.base_style = a   # цикл
        p = doc.add_paragraph()
        p.style = a
        # Смешанный регистр — по той же причине, что в TestStyleCapsCancellation:
        # изолируем наследование капса от регистровой эвристики этапа 2b.
        p.add_run("иванов Иван Иванович")
        doc.save(str(path))

        d = extract(str(path))   # не должно зависнуть
        # Капса в цикле нет -> detection_text не создан, но извлечение прошло.
        assert "detection_text" not in _seg(d, "p0").metadata
        assert _seg(d, "p0").text == "иванов Иван Иванович"


class TestStyleCapsRoundTripAndLength:
    def test_paragraph_style_caps_round_trip_char_for_char(self, tmp_path, config_path, ner_detector_module):
        """encrypt -> decrypt на .docx со СТИЛЕВЫМ капсом -> восстановленный текст
        посимвольно равен оригиналу (нормализованная версия в сессию не попадает)."""
        path = tmp_path / "d.docx"
        doc = Document()
        st = _add_para_style(doc, "CapsPara", all_caps=True)
        p = doc.add_paragraph()
        p.style = st
        p.add_run("стороны: иванов иван иванович")
        doc.save(str(path))

        d = extract(str(path))
        entities = detect_regex(d, config_path) + ner_detector_module.detect_ner(d, config_path)
        anon, result = tokenize(d, entities, config_path)
        session_id = save_session(result, storage_dir=str(tmp_path / "sessions"))
        restored, unresolved = detokenize(anon, session_id, storage_dir=str(tmp_path / "sessions"))

        assert unresolved == []
        assert restored == build_plain_text(d)
        assert "иванов иван иванович" in restored

    def test_length_invariant_for_style_caps(self, tmp_path):
        """len(detection_text) == len(text) для стилевого капса."""
        path = tmp_path / "d.docx"
        doc = Document()
        st = _add_para_style(doc, "CapsPara", all_caps=True)
        for txt in ("иванов иван иванович", "ип пирогова а.с."):
            p = doc.add_paragraph()
            p.style = st
            p.add_run(txt)
        doc.save(str(path))

        d = extract(str(path))
        created = 0
        for seg in d.segments:
            if "detection_text" in seg.metadata:
                created += 1
                assert len(seg.metadata["detection_text"]) == len(seg.text)
        assert created >= 2
