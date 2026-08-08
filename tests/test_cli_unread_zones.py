# -*- coding: utf-8 -*-
"""Этап 1b — тесты политики отказа на непрочитанных зонах (shifrator.py encrypt).

Через subprocess (реальный процесс, реальный exit code) — как и остальные тесты
блока 7. Изоляция хранилища: USERPROFILE/HOME подменяются на tmp_path/home, чтобы
CLI не трогал ~/.shifrator реального пользователя.

Три обязательных случая ЗАДАНИЯ:
  (а) документ с непрочитанной зоной, strict -> exit != 0, {sid}.txt НЕ создан;
  (б) он же с --allow-lossy -> обработан, sidecar создан и содержит текст зоны;
  (в) документ без зон, strict -> обработан как раньше (регресс не сломан).

ЭТАП NODES: носителем зоны здесь была ПДн в КОЛОНТИТУЛЕ. Колонтитулы теперь
читаются (часть 2 этапа), поэтому зоной они больше не являются и на роль
носителя не годятся. Носитель заменён на ВЛОЖЕННУЮ ТАБЛИЦУ — она осталась
непрочитанной зоной, и политика отказа проверяется на ней. Что колонтитул
теперь читается и маскируется — отдельный класс TestHeaderIsReadNow.
"""
import json
import os
import subprocess
import sys

import pytest

from testlib_policy import pin_all_types
from testlib_store import read_anon_text, read_unread_sidecar
from docx import Document

SHIFRATOR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "shifrator.py"
)

# ПДн, которые кладём в зону: они не должны попасть ни в анонимный текст, ни
# потеряться бесследно — sidecar обязан их показать.
ZONE_PII = "Смирнов Пётр Иванович, ИНН 7707083893"


def _run(args, tmp_path):
    env = dict(os.environ, PYTHONIOENCODING="utf-8",
               USERPROFILE=str(tmp_path / "home"), HOME=str(tmp_path / "home"))
    pin_all_types(tmp_path / "home")   # ЭТАП T1: набор «Максимум», см. testlib_policy
    return subprocess.run(
        [sys.executable, SHIFRATOR] + args,
        capture_output=True, text=True, encoding="utf-8",
        env=env, cwd=os.path.dirname(SHIFRATOR),
    )


def _sessions_dir(tmp_path):
    return tmp_path / "home" / ".shifrator" / "sessions"


@pytest.fixture
def docx_with_header(tmp_path):
    """Документ с ПДн во ВЛОЖЕННОЙ таблице — зона kind=nested_table.

    Имя фикстуры оставлено историческим: набор проверок политики отказа не
    зависит от вида зоны, менялся только её носитель (см. docstring модуля).
    """
    path = tmp_path / "with_zone.docx"
    document = Document()
    document.add_paragraph("Договор оказания услуг. Предмет договора — консультации.")
    outer = document.add_table(rows=1, cols=1)
    inner = outer.rows[0].cells[0].add_table(rows=1, cols=1)
    inner.rows[0].cells[0].text = ZONE_PII
    document.save(path)
    return path


@pytest.fixture
def docx_clean(tmp_path):
    """Документ без единой зоны — python-docx кладёт пустой header1.xml, и это
    НЕ должно приводить к отказу (пустой колонтитул не зона)."""
    path = tmp_path / "clean.docx"
    document = Document()
    document.add_paragraph("Договор оказания услуг. ИНН 7707083893.")
    document.save(path)
    return path


class TestStrictRefusal:
    """(а) strict — отказ, и ни следа сессии на диске."""

    def test_document_with_header_is_refused_with_nonzero_exit(self, docx_with_header, tmp_path):
        r = _run(["encrypt", str(docx_with_header)], tmp_path)
        assert r.returncode != 0, f"strict обязан отказать; stdout={r.stdout!r}"

    def test_refusal_prints_zone_table_and_lossy_hint_to_stdout(self, docx_with_header, tmp_path):
        r = _run(["encrypt", str(docx_with_header)], tmp_path)
        assert "nested_table" in r.stdout
        assert "word/document.xml" in r.stdout
        assert "--allow-lossy" in r.stdout

    def test_refusal_creates_no_session_txt(self, docx_with_header, tmp_path):
        _run(["encrypt", str(docx_with_header)], tmp_path)
        sessions = _sessions_dir(tmp_path)
        txts = list(sessions.glob("*.txt")) if sessions.exists() else []
        assert txts == [], f"отказ не должен оставлять {{sid}}.txt, найдено: {txts}"

    def test_refusal_does_not_print_session_id(self, docx_with_header, tmp_path):
        """stdout отказа не должен выглядеть как успех: никакого uuid, который
        пользователь (или скрипт) примет за session_id."""
        r = _run(["encrypt", str(docx_with_header)], tmp_path)
        import re
        assert not re.search(r"^[0-9a-f-]{36}$", r.stdout, re.M)


class TestTypedError:
    """Отказ — типизированное исключение, несущее сами зоны, а не только exit code.

    Проверяется на уровне функции (не subprocess): библиотечный потребитель должен
    получать UnreadZoneError с данными, а не sys.exit в лицо.
    """

    def test_check_unread_zones_raises_typed_error_carrying_zones(self, docx_with_header):
        sys.path.insert(0, os.path.dirname(SHIFRATOR))
        sys.path.insert(0, os.path.join(os.path.dirname(SHIFRATOR), "src"))
        import shifrator
        from unread_zones import UnreadZoneError

        with pytest.raises(UnreadZoneError) as exc:
            shifrator._check_unread_zones(str(docx_with_header), allow_lossy=False)

        assert [z.kind for z in exc.value.zones] == ["nested_table"]
        assert exc.value.path == str(docx_with_header)

    def test_check_unread_zones_returns_zones_in_lossy_mode(self, docx_with_header, capsys):
        sys.path.insert(0, os.path.dirname(SHIFRATOR))
        sys.path.insert(0, os.path.join(os.path.dirname(SHIFRATOR), "src"))
        import shifrator

        zones = shifrator._check_unread_zones(str(docx_with_header), allow_lossy=True)
        assert [z.kind for z in zones] == ["nested_table"]


class TestAllowLossy:
    """(б) --allow-lossy — работаем, но громко и с sidecar."""

    def test_allow_lossy_processes_document(self, docx_with_header, tmp_path):
        r = _run(["encrypt", str(docx_with_header), "--allow-lossy"], tmp_path)
        assert r.returncode == 0, f"stderr={r.stderr!r}"
        sid = r.stdout.strip().splitlines()[-1]
        assert (_sessions_dir(tmp_path) / f"{sid}.txt").is_file()

    def test_allow_lossy_writes_sidecar_with_raw_zone_text(self, docx_with_header, tmp_path):
        r = _run(["encrypt", str(docx_with_header), "--allow-lossy"], tmp_path)
        sid = r.stdout.strip().splitlines()[-1]
        sidecar = _sessions_dir(tmp_path) / f"{sid}.unread.json"
        assert sidecar.is_file(), "sidecar {sid}.unread.json не создан"

        # STORE: сайдкар зашифрован — внутри СЫРОЙ текст зон, то есть чистые ПДн.
        data = read_unread_sidecar(_sessions_dir(tmp_path), sid)
        assert data["lossy"] is True
        assert data["session_id"] == sid
        kinds = [z["kind"] for z in data["zones"]]
        assert "nested_table" in kinds
        # СЫРОЙ текст зоны — пользователь должен видеть, что именно выброшено.
        assert any(ZONE_PII in z["text"] for z in data["zones"])

    def test_allow_lossy_warns_loudly_on_stderr(self, docx_with_header, tmp_path):
        r = _run(["encrypt", str(docx_with_header), "--allow-lossy"], tmp_path)
        assert "ПРЕДУПРЕЖДЕНИЕ" in r.stderr
        assert "nested_table" in r.stderr

    def test_zone_text_is_absent_from_anonymized_output(self, docx_with_header, tmp_path):
        """Ключевой факт этапа: зона НЕ читается. Текст зоны не попадает в {sid}.txt
        ни открытым текстом, ни токеном — он просто отсутствует. Sidecar —
        единственное место, где пользователь его увидит."""
        r = _run(["encrypt", str(docx_with_header), "--allow-lossy"], tmp_path)
        sid = r.stdout.strip().splitlines()[-1]
        anon = read_anon_text(_sessions_dir(tmp_path), sid)   # STORE: файл зашифрован
        assert "Смирнов" not in anon


class TestConfigKey:
    """strict_zones: false в конфиге — то же, что --allow-lossy."""

    def test_strict_zones_false_in_config_allows_processing(self, docx_with_header, tmp_path):
        import shutil
        import yaml

        src_cfg = os.path.join(os.path.dirname(SHIFRATOR), "entity_types.yaml")
        cfg = tmp_path / "lossy.yaml"
        shutil.copy(src_cfg, cfg)
        data = yaml.safe_load(cfg.read_text(encoding="utf-8"))
        data["strict_zones"] = False
        cfg.write_text(yaml.safe_dump(data, allow_unicode=True), encoding="utf-8")

        r = _run(["encrypt", str(docx_with_header), "--config", str(cfg)], tmp_path)
        assert r.returncode == 0, f"stdout={r.stdout!r} stderr={r.stderr!r}"


class TestCleanDocumentNoRegression:
    """(в) документ без зон — работает ровно как раньше, strict ему не мешает."""

    def test_clean_document_processed_in_strict_mode(self, docx_clean, tmp_path):
        r = _run(["encrypt", str(docx_clean)], tmp_path)
        assert r.returncode == 0, f"stdout={r.stdout!r} stderr={r.stderr!r}"
        sid = r.stdout.strip()
        anon = read_anon_text(_sessions_dir(tmp_path), sid)   # STORE: файл зашифрован
        assert "[INN_1]" in anon

    def test_clean_document_gets_no_sidecar(self, docx_clean, tmp_path):
        r = _run(["encrypt", str(docx_clean)], tmp_path)
        sid = r.stdout.strip()
        assert not (_sessions_dir(tmp_path) / f"{sid}.unread.json").exists()

    def test_txt_input_unaffected_by_zone_policy(self, tmp_path):
        """.txt читается целиком — политика зон к нему неприменима и не должна
        добавить ни отказа, ни лишнего вывода в stdout."""
        txt = tmp_path / "input.txt"
        txt.write_text("ИНН 7707083893\n", encoding="utf-8")
        r = _run(["encrypt", str(txt)], tmp_path)
        assert r.returncode == 0
        assert len(r.stdout.strip().splitlines()) == 1


class TestHeaderIsReadNow:
    """ЭТАП NODES, часть 2: колонтитул — не зона, а обычный текст.

    До этого этапа тот же документ приводил к ОТКАЗУ (kind=header), а его ПДн
    не попадали в результат вовсе. Проверяем оба следствия смены контракта:
    отказа нет и ПДн замаскированы, то есть содержимое прошло тем же конвейером,
    что тело.
    """

    @pytest.fixture
    def docx_header_pii(self, tmp_path):
        path = tmp_path / "hdr.docx"
        document = Document()
        document.add_paragraph("Договор оказания услуг. Предмет — консультации.")
        document.sections[0].header.paragraphs[0].text = ZONE_PII
        document.sections[0].footer.paragraphs[0].text = "Телефон: +7 900 111-00-33"
        document.save(path)
        return path

    def test_header_document_is_no_longer_refused(self, docx_header_pii, tmp_path):
        r = _run(["encrypt", str(docx_header_pii)], tmp_path)
        assert r.returncode == 0, f"stdout={r.stdout!r} stderr={r.stderr!r}"

    def test_header_document_gets_no_sidecar(self, docx_header_pii, tmp_path):
        r = _run(["encrypt", str(docx_header_pii)], tmp_path)
        sid = r.stdout.strip()
        assert not (_sessions_dir(tmp_path) / f"{sid}.unread.json").exists()

    def test_header_and_footer_pii_are_masked(self, docx_header_pii, tmp_path):
        r = _run(["encrypt", str(docx_header_pii)], tmp_path)
        sid = r.stdout.strip()
        anon = read_anon_text(_sessions_dir(tmp_path), sid)
        assert "Смирнов" not in anon, f"ФИО из колонтитула утекло: {anon!r}"
        assert "7707083893" not in anon, f"ИНН из колонтитула утёк: {anon!r}"
        assert "111-00-33" not in anon, f"телефон из нижнего колонтитула утёк: {anon!r}"
        # содержимое дошло до результата, а не просто исчезло. Тип маски здесь
        # не проверяем: «Смирнов Пётр Иванович, ИНН 7707083893» — это ИНН-10,
        # то есть якорь ОРГАНИЗАЦИИ, и связка уезжает в ORG. Предмет проверки —
        # что колонтитул прошёл конвейер, а не какой слой его разметил.
        assert "[PHONE_" in anon
        assert anon.count("[") >= 3, f"колонтитул не размечен масками: {anon!r}"
