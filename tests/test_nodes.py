# -*- coding: utf-8 -*-
"""Этап NODES — программа либо читает зону, либо честно отказывает.

Молчаливый пропуск запрещён как класс. Тесты закрывают четыре части этапа:

  часть 1 — перепись узлов: наборы закрытого списка не пересекаются и лежат в
            ОДНОМ месте (двух копий у читателя и сторожа быть не может);
  часть 3 — четыре вида узлов (`w:ins`, `w:fldSimple`, `w:smartTag`, `w:sdt`
            обоих видов) читаются поимённо на `tests/fixtures/node_kinds.docx`;
  часть 4 — страж закрытого списка: узел вне переписи даёт ОТКАЗ, а не пропуск.

Фикстуры собраны генераторами `tests/fixtures/make_node_kinds_docx.py` и
`make_zone_kinds_docx.py`; данные полностью выдуманы.
"""
import os
import sys

import pytest

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, os.path.join(ROOT, "src"))

from docx import Document  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402

from extractor import extract  # noqa: E402
import unread_zones as uz  # noqa: E402
from unread_zones import scan_unread_zones  # noqa: E402

FIXTURES = os.path.join(ROOT, "tests", "fixtures")
NODE_KINDS = os.path.join(FIXTURES, "node_kinds.docx")
HEADERS = os.path.join(FIXTURES, "headers_footers.docx")
UNKNOWN = os.path.join(FIXTURES, "unknown_node.docx")


def _all_text(path: str) -> str:
    return "\n".join(s.text for s in extract(path).segments)


# --------------------------------------------------------------------------
# Часть 1 — перепись


class TestCensusRegistry:
    """Перепись — закрытый список, и он один на систему."""

    def test_reader_and_scanner_share_one_registry(self):
        """`extractor` не держит своей копии наборов: он импортирует их из
        `unread_zones`. Разойтись читателю и сторожу негде — а именно из-за
        расхождения и жил `DOG-NODE-SILENT`."""
        import extractor

        assert extractor.INLINE_DESCEND is uz.INLINE_DESCEND
        assert extractor.BLOCK_DESCEND is uz.BLOCK_DESCEND

    @pytest.mark.parametrize("descend,ignore", [
        ("INLINE_DESCEND", "INLINE_IGNORE"),
        ("BLOCK_DESCEND", "BLOCK_IGNORE"),
    ])
    def test_descend_and_ignore_do_not_overlap(self, descend, ignore):
        """Узел не может быть одновременно «спускаемся внутрь» и «пропускаем»:
        пересечение означало бы, что вердикт переписи зависит от порядка
        проверок в коде, а не от таблицы."""
        assert not (getattr(uz, descend) & getattr(uz, ignore))

    def test_four_node_kinds_are_in_the_registry(self):
        """Четыре вида узлов части 3 — именно ЧИТАЕМЫЕ, а не пропускаемые."""
        for name in ("ins", "fldSimple", "smartTag", "sdt"):
            tag = uz._W + name
            assert tag in uz.INLINE_DESCEND, name
            assert tag not in uz.INLINE_IGNORE, name

    def test_deleted_revision_text_is_ignored_deliberately(self):
        """`w:del` / `w:moveFrom` — удалённый текст рецензирования: НЕ видимое
        содержимое, вердикт переписи «пропускаем осознанно». Его отсутствие в
        наборе игнора превратило бы его в unknown_node и в ложный отказ."""
        assert uz._W + "del" in uz.INLINE_IGNORE
        assert uz._W + "moveFrom" in uz.INLINE_IGNORE


# --------------------------------------------------------------------------
# Часть 3 — четыре вида узлов, поимённо

# Маркер узла -> опознавательный телефон из фикстуры (см. make_node_kinds_docx).
READ_CASES = [
    ("обычный ран (контроль)", "+7 900 000-00-01"),
    ("w:ins — вставка рецензирования", "+7 900 000-00-02"),
    ("w:hyperlink — гиперссылка", "+7 900 000-00-04"),
    ("w:fldSimple — результат простого поля", "+7 900 000-00-05"),
    ("w:smartTag — смарт-тег", "+7 900 000-00-06"),
    ("w:sdt inline — элемент управления в строке", "+7 900 000-00-07"),
    ("w:sdt блочный — элемент управления верхнего уровня", "+7 900 000-00-08"),
]


class TestFourNodeKinds:
    """Каждый вид узла назван, и по своему телефону видно, прочитан ли он."""

    @pytest.mark.parametrize("what,phone", READ_CASES)
    def test_node_content_is_read(self, what, phone):
        assert phone in _all_text(NODE_KINDS), (
            f"{what}: содержимое узла не дошло до сегментов — это тихая потеря "
            f"(и раскол якоря, если в узле лежал маркер типа)"
        )

    def test_deleted_text_is_not_read(self):
        """`w:del` — единственный из видов узлов фикстуры, чей текст читать
        НЕЛЬЗЯ: удалённая правка не является видимым содержимым документа."""
        assert "+7 900 000-00-03" not in _all_text(NODE_KINDS)

    def test_no_silent_zone_left_on_the_fixture(self):
        """Раз всё прочитано — объявлять нечего. Пустой список зон здесь
        осмыслен только вместе с проверками выше."""
        assert scan_unread_zones(NODE_KINDS) == []

    def test_anchor_split_is_closed(self):
        """`DOG-ANCHOR-SPLIT`: якорь в обёртке, значение — обычным раном.

        До этапа NODES обёртка прятала якорь, якорный детектор маркера не
        видел, и ИНН уходил в анонимный текст ОТКРЫТЫМ. Проверяем не маску, а
        причину: оба куска абзаца обязаны оказаться в ОДНОМ сегменте, вплотную.
        """
        import tempfile

        document = Document()
        p = document.add_paragraph()
        ins = OxmlElement("w:ins")
        ins.set(uz._W + "id", "1")
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = "ИНН "
        r.append(t)
        ins.append(r)
        p._p.append(ins)
        p.add_run("502707033944")

        with tempfile.TemporaryDirectory() as d:
            path = os.path.join(d, "split.docx")
            document.save(path)
            texts = [s.text for s in extract(path).segments]

        assert "ИНН 502707033944" in texts, (
            f"якорь и значение не собрались в один сегмент: {texts!r}")


# --------------------------------------------------------------------------
# Часть 2 — колонтитулы


class TestHeadersAndFooters:
    """Колонтитул читается тем же конвейером, что тело, и знает свою часть."""

    def test_header_and_footer_text_is_extracted(self):
        text = _all_text(HEADERS)
        assert "Ковалёва Нина Петровна" in text
        assert "+7 900 111-00-11" in text

    def test_footer_table_is_extracted(self):
        text = _all_text(HEADERS)
        assert "+7 900 111-00-22" in text, "таблица в колонтитуле не прочитана"

    def test_segments_carry_their_part_name(self):
        parts = {s.metadata.get("part") for s in extract(HEADERS).segments}
        assert "word/header1.xml" in parts
        assert "word/footer1.xml" in parts
        assert None in parts, "у сегментов тела ключа part быть не должно"

    def test_table_index_is_unique_across_parts(self):
        """`table_index` — ключ группировки ячеек при сборке текста и признак
        «одна таблица» у граничного окна детекции. Совпади номер у таблицы тела
        и таблицы колонтитула — они склеились бы в одну."""
        segs = extract(HEADERS).segments
        seen = {}
        for s in segs:
            if s.source_type != "docx_table_cell":
                continue
            idx = s.metadata["table_index"]
            part = s.metadata.get("part")
            assert seen.setdefault(idx, part) == part, (
                f"table_index={idx} встретился в двух разных частях")

    def test_headers_are_no_longer_declared_a_zone(self):
        assert scan_unread_zones(HEADERS) == []

    def test_part_order_is_deterministic(self):
        first = [s.id for s in extract(HEADERS).segments]
        second = [s.id for s in extract(HEADERS).segments]
        assert first == second
        assert first.index("header1:p0") < first.index("footer1:p0"), (
            "верхний колонтитул обязан идти перед нижним")


# --------------------------------------------------------------------------
# Часть 4 — страж закрытого списка


class TestClosedListGuard:
    """Живая проба: узел, которого нет в переписи, роняет чтение отказом."""

    def test_unknown_node_is_declared_a_zone(self):
        zones = scan_unread_zones(UNKNOWN)
        kinds = [z.kind for z in zones]
        assert kinds == ["unknown_node"], f"ожидался отказной узел, получено {zones!r}"

    def test_zone_names_the_offending_tag(self):
        """Пользователь обязан узнать, ЧТО именно система не поняла, — иначе
        отказ неотличим от поломки."""
        (zone,) = scan_unread_zones(UNKNOWN)
        assert zone.detail == "m:oMath"
        assert "m:oMath" in zone.text_preview

    def test_unknown_node_is_not_silently_read(self):
        """Содержимое незнакомого узла не должно просочиться в сегменты как
        якобы прочитанное: отказ и тихое чтение — разные вещи."""
        assert "7736050003" not in _all_text(UNKNOWN)

    def test_unknown_node_in_header_is_also_caught(self):
        """Сторож работает во ВСЕХ читаемых частях, не только в теле: иначе
        колонтитул, которому этап NODES открыл дорогу, стал бы новой дырой."""
        import tempfile

        document = Document()
        document.add_paragraph("Тело.")
        hp = document.sections[0].header.paragraphs[0]
        hp.text = "Колонтитул: "
        # ЭТАП CLIENT: у формулы теперь есть СОДЕРЖИМОЕ. До этапа здесь висел
        # пустой `m:oMath`, и проба проходила только потому, что отказ шёл на
        # любом узле вне переписи — даже на заведомо пустом. Часть 2 этапа сузила
        # отказ до узлов, внутри которых есть что терять, поэтому фикстура
        # приведена к тому же виду, что и дисковая (`_omath("ИНН …")`): проба
        # проверяет «сторож работает и в колонтитуле», а не «пустое отказывает».
        omath = OxmlElement("m:oMath")
        r = OxmlElement("m:r")
        t = OxmlElement("m:t")
        t.text = "ИНН 7736050003"
        r.append(t)
        omath.append(r)
        hp._p.append(omath)

        with tempfile.TemporaryDirectory() as d:
            path = os.path.join(d, "hdr_unknown.docx")
            document.save(path)
            zones = scan_unread_zones(path)

        assert [(z.kind, z.part) for z in zones] == [
            ("unknown_node", "word/header1.xml")]

    def test_wrapper_around_table_row_is_caught(self):
        """Обёртка, которую ЧИТАТЕЛЬ не видит, обязана быть отказом, даже если
        сама по себе она законный OOXML. python-docx перечисляет только прямые
        `w:tr` таблицы, поэтому `w:sdt` вокруг строки прячет её целиком."""
        import tempfile

        document = Document()
        table = document.add_table(rows=1, cols=1)
        table.rows[0].cells[0].text = "ИНН 7736050003"
        tr = table.rows[0]._tr
        tbl = tr.getparent()
        sdt = OxmlElement("w:sdt")
        content = OxmlElement("w:sdtContent")
        tbl.replace(tr, sdt)
        sdt.append(content)
        content.append(tr)

        with tempfile.TemporaryDirectory() as d:
            path = os.path.join(d, "wrapped_row.docx")
            document.save(path)
            zones = scan_unread_zones(path)
            text = _all_text(path)

        assert [z.kind for z in zones] == ["unknown_node"]
        assert "7736050003" not in text


# --------------------------------------------------------------------------
# ЭТАП CLIENT, часть 2 — сужение отказа на узле вне переписи.
#
# `mc:AlternateContent` (новый Word, LibreOffice: фигуры, надписи, SmartArt) до
# этапа отказывал ВСЕГДА — в том числе на фигуре без единого символа. Отказ
# сужен ровно до случаев, где есть что терять. Молчаливый пропуск не заведён
# ни в одном из них: каждая проба ниже либо отказывает, либо доказывает, что
# терять нечего.

_ALT_NS = (
    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
    'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" '
    'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" '
    'xmlns:v="urn:schemas-microsoft-com:vml" '
    'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
)


def _alt(choice_inner: str, fallback_inner: str) -> str:
    return (f"<mc:AlternateContent {_ALT_NS}>"
            f'<mc:Choice Requires="wps">{choice_inner}</mc:Choice>'
            f"<mc:Fallback>{fallback_inner}</mc:Fallback>"
            f"</mc:AlternateContent>")


def _txbx(tag: str, text: str) -> str:
    return (f"<{tag}><w:txbxContent><w:p><w:r><w:t>{text}</w:t></w:r></w:p>"
            f"</w:txbxContent></{tag}>")


def _docx_with(tmp_path, xml: str, name: str = "alt.docx") -> str:
    """Документ с одним абзацем текста и приложенным к нему сырым узлом."""
    from docx.oxml import parse_xml

    document = Document()
    p = document.add_paragraph("Тело договора.")
    p._p.append(parse_xml(xml))
    path = os.path.join(str(tmp_path), name)
    document.save(path)
    return path


class TestAlternateContentNarrowing:

    def test_shape_with_text_still_refuses(self, tmp_path):
        """Надпись с текстом — отказ как прежде. Это главный инвариант: сужение
        не имеет права выпустить наружу ни одного значения."""
        path = _docx_with(tmp_path, _alt(_txbx("wps:txbx", "ИНН 7736050003"),
                                         _txbx("v:textbox", "ИНН 7736050003")))
        zones = scan_unread_zones(path)
        assert [z.kind for z in zones] == ["unknown_node"]
        assert zones[0].detail == "mc:AlternateContent"
        assert "7736050003" not in _all_text(path)

    def test_two_branches_are_not_counted_twice(self, tmp_path):
        """`mc:Choice` и `mc:Fallback` — ДВА кодирования одного содержимого.
        Наивная склейка показывала пользователю текст дважды (72 символа там,
        где текста 36) и столько же писала в sidecar lossy-режима."""
        text = "ИНН 7736050003"
        path = _docx_with(tmp_path, _alt(_txbx("wps:txbx", text),
                                         _txbx("v:textbox", text)))
        (zone,) = scan_unread_zones(path)
        assert zone.char_count == len(text)
        assert zone.text == text

    def test_text_only_in_fallback_still_refuses(self, tmp_path):
        """Предпочтение `mc:Choice` не имеет права ослепить нас к тексту,
        который лежит ТОЛЬКО в запасной ветви."""
        path = _docx_with(tmp_path, _alt("<wps:spPr/>",
                                         _txbx("v:textbox", "ИНН 7736050003")))
        (zone,) = scan_unread_zones(path)
        assert zone.text == "ИНН 7736050003"

    def test_shape_without_any_text_does_not_refuse(self, tmp_path):
        """Фигура без единого символа (линия, рамка) — терять нечего, и отказ
        на ней был отказом на пустом месте. Ровно этот случай этап и снимает."""
        path = _docx_with(tmp_path, _alt("<wps:spPr/>",
                                         '<w:pict><v:shape style="width:9pt"/></w:pict>'))
        assert scan_unread_zones(path) == []

    def test_text_outside_w_t_is_still_seen(self, tmp_path):
        """Проверка пустоты идёт ШИРЕ, чем `w:t`. Формула держит текст в `m:t`,
        и узкая проверка превратила бы отказ в тихий пропуск — на дисковой
        фикстуре `unknown_node.docx` её `m:oMath` по `w:t` даёт РОВНО НОЛЬ."""
        from docx.oxml import parse_xml

        omath = parse_xml(
            '<m:oMath xmlns:m="http://schemas.openxmlformats.org/'
            'officeDocument/2006/math"><m:r><m:t>7736050003</m:t></m:r></m:oMath>')
        assert uz._zone_text(omath) == ""          # узкая проверка слепа
        assert uz._subtree_has_text(omath) is True  # широкая — видит
        assert [z.kind for z in scan_unread_zones(UNKNOWN)] == ["unknown_node"]

    def test_wordart_text_in_attribute_is_seen(self, tmp_path):
        """WordArt держит надпись в АТРИБУТЕ `v:textpath/@string`; `itertext()`
        до неё не добирается."""
        path = _docx_with(tmp_path, _alt(
            "<wps:spPr/>",
            '<w:pict><v:shape><v:textpath string="Орлов Сергей Ильич"/>'
            "</v:shape></w:pict>"))
        assert [z.kind for z in scan_unread_zones(path)] == ["unknown_node"]

    def test_reference_to_non_image_part_still_refuses(self, tmp_path):
        """Ссылка на НЕрастровую часть (диаграмма, OLE, SmartArt) означает
        содержимое в другой части пакета, которую мы не смотрели. Символов в
        самом узле ноль — и всё равно отказ."""
        from docx.oxml import parse_xml

        document = Document()
        p = document.add_paragraph("Тело договора.")
        # Связь заводим настоящую, средствами пакета (иначе python-docx не
        # откроет файл). Цель связи для проверки безразлична — важен ТИП.
        rid = document.part.relate_to(
            document.part, "http://schemas.openxmlformats.org/"
            "officeDocument/2006/relationships/chart")
        p._p.append(parse_xml(_alt(f'<wps:spPr><c:chart xmlns:c="http://'
                                   f'schemas.openxmlformats.org/drawingml/2006/chart" '
                                   f'r:id="{rid}"/></wps:spPr>',
                                   '<w:pict><v:shape style="width:9pt"/></w:pict>')))
        path = os.path.join(str(tmp_path), "chart.docx")
        document.save(path)
        assert [z.kind for z in scan_unread_zones(path)] == ["unknown_node"]

    def test_reference_to_image_part_does_not_refuse(self, tmp_path):
        """Растр текста не несёт, и обычная картинка в теле отказа не вызывает
        и сегодня. Обёртка вокруг неё не имеет права менять вердикт."""
        from docx.oxml import parse_xml

        document = Document()
        p = document.add_paragraph("Тело договора.")
        rid = document.part.relate_to(
            document.part, "http://schemas.openxmlformats.org/"
            "officeDocument/2006/relationships/image")
        p._p.append(parse_xml(_alt(f'<wps:spPr><a:blip xmlns:a="http://'
                                   f'schemas.openxmlformats.org/drawingml/2006/main" '
                                   f'r:embed="{rid}"/></wps:spPr>',
                                   '<w:pict><v:shape style="width:9pt"/></w:pict>')))
        path = os.path.join(str(tmp_path), "img.docx")
        document.save(path)
        assert scan_unread_zones(path) == []

    def test_unresolvable_reference_refuses(self, tmp_path):
        """Ссылка, которую нечем разрешить (нет .rels, битый .rels, чужой rId),
        обязана оставить отказ: положительный ответ «это картинка» мы обязаны
        ПОЛУЧИТЬ, а не предположить."""
        path = _docx_with(tmp_path, _alt(
            '<wps:spPr><a:blip xmlns:a="http://schemas.openxmlformats.org/'
            'drawingml/2006/main" r:embed="rIdНеСуществует"/></wps:spPr>',
            '<w:pict><v:shape style="width:9pt"/></w:pict>'))
        assert [z.kind for z in scan_unread_zones(path)] == ["unknown_node"]
