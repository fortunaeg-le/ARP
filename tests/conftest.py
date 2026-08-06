"""Общие фикстуры для тестов SHIFRATOR.

Тяжёлые импорты (natasha через ner_detector) грузятся один раз за сессию
через module/session-scope фикстуры, чтобы не платить ~1с и сотни МБ RAM
на каждый тестовый модуль отдельно (см. HANDOFF_3, "Побочные эффекты импорта").
"""
import os
import sys
from pathlib import Path

import pytest
from docx import Document

_PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, _PROJECT_ROOT)

CONFIG_PATH = os.path.join(_PROJECT_ROOT, "entity_types.yaml")


@pytest.fixture(scope="session", autouse=True)
def isolated_home(tmp_path_factory):
    """ЭТАП STORE. Домашний каталог на ВЕСЬ прогон тестов уводится во временный.

    Продукт хранит сессии, разметку и настройки в `~/.shifrator`. Дважды за
    последние сессии тест писал туда РЕАЛЬНОМУ пользователю — оба раза чинилось
    точечной заплаткой в конкретном тесте, и оба раза следующий новый тест про
    заплатку не знал. Точечный приём и не мог сработать как правило: он требует,
    чтобы автор теста ВСПОМНИЛ про ловушку.

    Здесь общий сторож: HOME/USERPROFILE подменены до первого теста, и
    `Path.home()` (а значит и `default_storage_dir()`, и `type_policy`) во всех
    тестах смотрит во временный каталог. HOMEDRIVE/HOMEPATH снимаются: на Windows
    `expanduser` берёт их РАНЬШЕ USERPROFILE, и без снятия подмена не подействует.
    Подмена наследуется и запускаемыми из тестов subprocess'ами (они получают
    os.environ), поэтому CLI-тесты закрыты тем же ходом.

    Это ВТОРОЙ рубеж, а не замена починки: сам дефект (путь застывал на импорте)
    исправлен в `session_store.default_storage_dir()`. Без той починки эта
    фикстура не работала бы вовсе — модуль импортируется раньше неё.
    """
    home = tmp_path_factory.mktemp("home")
    with pytest.MonkeyPatch.context() as mp:
        mp.setenv("USERPROFILE", str(home))
        mp.setenv("HOME", str(home))
        mp.delenv("HOMEDRIVE", raising=False)
        mp.delenv("HOMEPATH", raising=False)
        assert Path.home() == home, (
            f"подмена домашнего каталога не подействовала: Path.home()={Path.home()}"
        )
        yield home


@pytest.fixture(scope="session")
def config_path():
    return CONFIG_PATH


@pytest.fixture(scope="session")
def ner_detector_module():
    """Импортирует ner_detector (и тем самым natasha) один раз за сессию тестов."""
    import ner_detector
    return ner_detector


def make_docx(path, paragraphs_before, table_rows, paragraphs_after, heading=None):
    """Строит .docx: опциональный heading, параграфы before, таблица NxM, параграфы after.

    table_rows: list[list[str]] — текст ячеек по строкам.
    Возвращает путь к файлу.
    """
    doc = Document()
    if heading is not None:
        doc.add_heading(heading, level=1)
    for p in paragraphs_before:
        doc.add_paragraph(p)
    if table_rows:
        n_rows = len(table_rows)
        n_cols = len(table_rows[0])
        table = doc.add_table(rows=n_rows, cols=n_cols)
        for r, row in enumerate(table_rows):
            for c, text in enumerate(row):
                table.cell(r, c).text = text
    for p in paragraphs_after:
        doc.add_paragraph(p)
    doc.save(path)
    return str(path)


@pytest.fixture
def docx_factory(tmp_path):
    """Фабрика .docx файлов в tmp_path. Использование:
    path = docx_factory("name.docx", paragraphs_before=[...], table_rows=[[...]], paragraphs_after=[...])
    """
    def _factory(filename="doc.docx", paragraphs_before=(), table_rows=None, paragraphs_after=(), heading=None):
        path = tmp_path / filename
        return make_docx(path, list(paragraphs_before), table_rows, list(paragraphs_after), heading=heading)
    return _factory


@pytest.fixture
def txt_factory(tmp_path):
    """Фабрика .txt файлов в tmp_path с заданным содержимым и кодировкой."""
    def _factory(content, filename="doc.txt", encoding="utf-8-sig"):
        path = tmp_path / filename
        with open(path, "w", encoding=encoding, newline="") as f:
            f.write(content)
        return str(path)
    return _factory
