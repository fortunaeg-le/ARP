"""Генератор фикстуры `node_kinds.docx` — по одному виду узла OOXML на абзац.

Зачем. `extractor._extract_docx` читает тело через `document.element.body.iterchildren()`
и берёт только `w:p` и `w:tbl`; текст абзаца берётся у python-docx как
`CT_P.text` = `"".join(e.text for e in self.xpath("w:r | w:hyperlink"))`, то есть
только ПРЯМЫЕ дети `w:p`. Всё, что заворачивает раны в промежуточный узел
(`w:ins`, `w:fldSimple`, `w:smartTag`, `w:sdt`), из текста выпадает. Сканер
непрочитанных зон (`unread_zones._scan_document`) ищет только `w:txbxContent` и
вложенные `w:tbl`, поэтому такие узлы не объявляются и зоной.

Фикстура даёт контролируемый случай: в каждом виде узла лежит СВОЙ порождённый
телефон (`+7 900 000-00-NN`) и СВОЁ порождённое ФИО. По тому, какие номера
дошли до детекции, видно поимённо, что прочитано, а что потеряно.

Данные полностью выдуманы. Реальных документов в репозитории нет и быть не может.

Запуск: venv/Scripts/python.exe tests/fixtures/make_node_kinds_docx.py
"""

from pathlib import Path

from docx import Document
from docx.oxml.ns import nsmap, qn
from docx.oxml import OxmlElement

# Маркер -> (номер телефона, ФИО). Номер служит опознавательным знаком зоны:
# по последним двум цифрам однозначно видно, из какого узла значение пришло.
CASES = [
    ("plain", "+7 900 000-00-01", "Иванов Пётр Сергеевич"),
    ("ins", "+7 900 000-00-02", "Петрова Ольга Ивановна"),
    ("del", "+7 900 000-00-03", "Сидоров Илья Львович"),
    ("hyperlink", "+7 900 000-00-04", "Кузнецова Мария Дмитриевна"),
    ("fldSimple", "+7 900 000-00-05", "Смирнов Артём Игоревич"),
    ("smartTag", "+7 900 000-00-06", "Ершова Татьяна Дмитриевна"),
    ("sdt_inline", "+7 900 000-00-07", "Крылов Никита Павлович"),
    ("sdt_block", "+7 900 000-00-08", "Волкова Анна Борисовна"),
]


def _run(text: str, tag: str = "w:t"):
    """`w:r` с одним текстовым узлом. tag='w:delText' — для удалённого текста."""
    r = OxmlElement("w:r")
    t = OxmlElement(tag)
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _payload(marker: str) -> str:
    phone, fio = dict((m, (p, f)) for m, p, f in CASES)[marker]
    return f"{fio}, тел. {phone}"


def _wrapped(marker: str, wrapper_tag: str, *, inner_tag: str = "w:t",
             attrs: dict | None = None, content_tag: str | None = None):
    """Узел `wrapper_tag`, внутри которого лежит ран с полезной нагрузкой.

    content_tag — промежуточный контейнер (у `w:sdt` это `w:sdtContent`).
    """
    node = OxmlElement(wrapper_tag)
    for k, v in (attrs or {}).items():
        node.set(qn(k), v)
    holder = node
    if content_tag is not None:
        holder = OxmlElement(content_tag)
        node.append(holder)
    holder.append(_run(_payload(marker), inner_tag))
    return node


def build(path: Path) -> Path:
    document = Document()

    # 0. Контроль: обычный ран прямым ребёнком w:p. Обязан читаться.
    document.add_paragraph("Контроль. " + _payload("plain"))

    # 1. w:ins — вставка в режиме рецензирования. Видимый текст документа.
    p = document.add_paragraph("Правка-вставка. ")
    p._p.append(_wrapped("ins", "w:ins",
                         attrs={"w:id": "101", "w:author": "tester",
                                "w:date": "2026-01-01T00:00:00Z"}))

    # 2. w:del — удаление в режиме рецензирования, текст лежит в w:delText.
    #    НЕ является видимым содержимым: правильное поведение — не выводить его.
    p = document.add_paragraph("Правка-удаление. ")
    p._p.append(_wrapped("del", "w:del", inner_tag="w:delText",
                         attrs={"w:id": "102", "w:author": "tester",
                                "w:date": "2026-01-01T00:00:00Z"}))

    # 3. w:hyperlink — гиперссылка. Видимый текст, r:id намеренно не ставим:
    #    предмет проверки — чтение текста, а не разрешение связи.
    p = document.add_paragraph("Гиперссылка. ")
    p._p.append(_wrapped("hyperlink", "w:hyperlink"))

    # 4. w:fldSimple — простое поле Word; внутри лежит РЕЗУЛЬТАТ поля.
    p = document.add_paragraph("Простое поле. ")
    p._p.append(_wrapped("fldSimple", "w:fldSimple",
                         attrs={"w:instr": " DOCPROPERTY Party "}))

    # 5. w:smartTag — распознанная сущность старых версий Word.
    p = document.add_paragraph("Смарт-тег. ")
    p._p.append(_wrapped("smartTag", "w:smartTag",
                         attrs={"w:element": "PersonName"}))

    # 6. w:sdt внутри абзаца (inline content control).
    p = document.add_paragraph("Поле формы в строке. ")
    p._p.append(_wrapped("sdt_inline", "w:sdt", content_tag="w:sdtContent"))

    # 7. w:sdt на ВЕРХНЕМ уровне тела (block-level content control): внутри
    #    w:sdtContent лежит целый w:p. Тело обходится iterchildren() и такой
    #    узел не является ни w:p, ни w:tbl.
    sdt = OxmlElement("w:sdt")
    content = OxmlElement("w:sdtContent")
    block_p = OxmlElement("w:p")
    block_p.append(_run("Блочное поле формы. " + _payload("sdt_block")))
    content.append(block_p)
    sdt.append(content)
    document.element.body.append(sdt)

    document.save(str(path))
    return path


if __name__ == "__main__":
    out = build(Path(__file__).with_name("node_kinds.docx"))
    print(f"записано: {out}")
    print("пространства имён:", sorted(nsmap))
