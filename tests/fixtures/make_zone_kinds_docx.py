"""Генератор фикстур этапа NODES: колонтитулы и узел вне переписи.

Две фикстуры, обе с полностью выдуманными данными (реальных документов в
репозитории нет и быть не может):

* `headers_footers.docx` — верхний и нижний колонтитулы с ПДн, плюс таблица в
  нижнем колонтитуле. До этапа NODES колонтитулы объявлялись непрочитанной
  зоной, и encrypt на таком документе отказывал; теперь они читаются тем же
  конвейером, что тело, а восстановление обязано попасть в СВОЮ часть архива
  (`word/header1.xml` / `word/footer1.xml`), а не в `word/document.xml`.

* `unknown_node.docx` — абзац, внутри которого лежит узел, которого НЕТ в
  переписи закрытого списка (`unread_zones`, наборы `*_DESCEND` / `*_IGNORE`).
  Берём `m:oMath` — настоящий узел OOXML (формула), намеренно не внесённый в
  перепись: у формулы текст лежит в `m:t`, конвейер её не читает и читать не
  умеет. Ожидаемое поведение — зона `unknown_node` и ОТКАЗ, а не тихий пропуск.

Запуск: venv/Scripts/python.exe tests/fixtures/make_zone_kinds_docx.py
"""

from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

_M = "http://schemas.openxmlformats.org/officeDocument/2006/math"

# Опознавательные значения: по номеру телефона видно, из какой части он пришёл.
HEADER_TEXT = "Договор № 7/2026. Исполнитель: Ковалёва Нина Петровна"
HEADER_TEXT_2 = "Исходный файл: dogovor_ivanov_2026.docx"
FOOTER_TEXT = "Контакт: +7 900 111-00-11, ИНН 7736050003"
FOOTER_CELLS = ("Стр. 1 из 3", "Телефон: +7 900 111-00-22")


def build_headers_footers(path: Path) -> Path:
    document = Document()
    document.add_paragraph("Тело договора. Заказчик: Орлов Сергей Ильич.")
    document.add_paragraph("Реквизиты в теле: ИНН 502707033944.")

    section = document.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    header.paragraphs[0].text = HEADER_TEXT
    header.add_paragraph(HEADER_TEXT_2)

    footer = section.footer
    footer.is_linked_to_previous = False
    footer.paragraphs[0].text = FOOTER_TEXT
    table = footer.add_table(rows=1, cols=2, width=Inches(3))
    table.rows[0].cells[0].text = FOOTER_CELLS[0]
    table.rows[0].cells[1].text = FOOTER_CELLS[1]

    document.save(str(path))
    return path


def _omath(text: str):
    """`m:oMath` с одним `m:r`/`m:t` — узел вне переписи закрытого списка."""
    omath = OxmlElement("m:oMath")
    r = OxmlElement("m:r")
    t = OxmlElement("m:t")
    t.text = text
    r.append(t)
    omath.append(r)
    return omath


def build_unknown_node(path: Path) -> Path:
    document = Document()
    document.add_paragraph("Обычный абзац. Заказчик: Орлов Сергей Ильич.")
    p = document.add_paragraph("Формула: ")
    p._p.append(_omath("ИНН 7736050003"))
    document.save(str(path))
    return path


if __name__ == "__main__":
    here = Path(__file__).parent
    print("записано:", build_headers_footers(here / "headers_footers.docx"))
    print("записано:", build_unknown_node(here / "unknown_node.docx"))
