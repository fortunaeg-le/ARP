# -*- coding: utf-8 -*-
import os, sys, io
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "src"))
from docx import Document
from models import SourceDocument, TextSegment
from regex_detector import detect_regex
from ner_detector import detect_ner
from syntax_compound import merge_compound_entities
from tokenizer import tokenize
from extractor import extract
CFG = os.path.join(os.path.dirname(__file__), "..", "entity_types.yaml")
OUT=io.StringIO()
def w(*a): OUT.write(" ".join(str(x) for x in a)+"\n")

def enc(doc):
    ents = detect_regex(doc, CFG)+detect_ner(doc, CFG)
    ents = merge_compound_entities(doc, ents)
    return tokenize(doc, ents, CFG)

w("=== F. Граничная сущность внутри all_caps-сегментов (волна1 x волна2) ===")
# ИНН, разорванный на два all_caps параграфа
d=Document()
for txt in ["инн 7707", "083893"]:
    p=d.add_paragraph(); r=p.add_run(txt); r.font.all_caps=True
p1=os.path.join(os.path.dirname(__file__),"f_inn.docx"); d.save(p1)
doc=extract(p1)
w("сегменты:", [(s.text, s.metadata.get('detection_text')) for s in doc.segments])
anon,final=enc(doc)
w("  ИНН-boundary anon:", repr(anon))
w("  утечка '7707083893' целиком:", "7707083893" in anon.replace("\n",""))
w("  токены:", [(e.entity_type,e.original_text) for e in final])

# ФИО, разорванное на два all_caps параграфа
d2=Document()
for txt in ["иванов иван", "иванович"]:
    p=d2.add_paragraph(); r=p.add_run(txt); r.font.all_caps=True
p2=os.path.join(os.path.dirname(__file__),"f_per.docx"); d2.save(p2)
doc2=extract(p2)
anon2,final2=enc(doc2)
w("\n  ФИО-boundary сегменты:", [(s.text, s.metadata.get('detection_text')) for s in doc2.segments])
w("  ФИО-boundary anon:", repr(anon2))
w("  токены:", [(e.entity_type,e.original_text) for e in final2])

with open(os.path.join(os.path.dirname(__file__),"f_out.txt"),"w",encoding="utf-8") as f:
    f.write(OUT.getvalue())
print("done")
