# -*- coding: utf-8 -*-
import os, sys, io, time
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "src"))
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from models import SourceDocument, TextSegment
from regex_detector import detect_regex
from ner_detector import detect_ner
from syntax_compound import merge_compound_entities
from tokenizer import tokenize, _detect_boundary_entities
from extractor import extract
CFG = os.path.join(os.path.dirname(__file__), "..", "entity_types.yaml")
OUT = io.StringIO()
def w(*a): OUT.write(" ".join(str(x) for x in a)+"\n")

# ================= B. ALL CAPS via PARAGRAPH STYLE inheritance =================
w("========== B. ALL CAPS (наследование от стиля абзаца) ==========")
d = Document()
st = d.styles.add_style("CapsPara", WD_STYLE_TYPE.PARAGRAPH)
st.font.all_caps = True
p = d.add_paragraph(style=st)
p.add_run("иванов иван иванович")          # lowercase in XML, caps via STYLE only
d.add_paragraph().add_run("инн 7707083893")
path = os.path.join(os.path.dirname(__file__), "caps_style.docx")
d.save(path)
doc = extract(path)
seg0 = doc.segments[0]
w("seg0.text      =", repr(seg0.text))
w("detection_text =", repr(seg0.metadata.get("detection_text")))
w("равные длины   =", len(seg0.text) == len(seg0.metadata.get("detection_text", seg0.text)))
ents = detect_regex(doc, CFG) + detect_ner(doc, CFG)
pers = [e for e in ents if e.entity_type == "PERSON"]
w("PERSON найден (через стиль, не прямой формат):", [e.original_text for e in pers])
# round-trip char-for-char
anon, final = tokenize(doc, ents, CFG)
from detokenizer import detokenize
# reconstruct mapping decrypt: replace tokens back
# simple: build token->original map
tokmap = {}
for e in final:
    pass
# Use session-like: original preserved in final entities; verify original_text slices match
ok = all(seg_by.text[e.start:e.end] == e.original_text
         for seg_by in doc.segments for e in final if e.segment_id == seg_by.id)
w("инвариант original_text==slice для all-caps сегмента:", ok)

# ================= C. Compound negatives (>=5) =================
w("\n========== C. Составные — НЕГАТИВНЫЕ (разные стороны НЕ склеивать) ==========")
NEG = [
    "Стороны: ИП Иванов И.И. и ООО «Ромашка»",
    "ИП Иванов И.И., ООО «Ромашка», ЗАО «Вектор»",
    "Договор между ИП Петров П.П. и ИП Сидоров С.С.",
    "Поставщик ООО «Заря»\nПокупатель Иванов Иван Иванович",
    "ИП Иванов И.И.\nООО «Ромашка»",
    "В список включены: Смирнов А.А. и ООО «Лидер»",
    "Глава КФХ Петров Пётр Петрович",              # role+ФИО не сливать
    "Генеральный директор Иванов И.И. ООО «Ромашка»",
]
def compound(text):
    seg = TextSegment(id="s0", text=text, source_type="txt_line", metadata={})
    doc = SourceDocument(segments=[seg], source_format="txt", source_path="<c>")
    ents = detect_regex(doc, CFG) + detect_ner(doc, CFG)
    m = merge_compound_entities(doc, ents)
    return m
for text in NEG:
    m = compound(text)
    orgs = [e.original_text for e in m if e.entity_type == "ORG"]
    bad = [o for o in orgs if ("Иванов" in o and "Ромашка" in o) or ("Петров" in o and "Сидоров" in o)]
    merged_two_parties = any((("Иванов" in o or "Смирнов" in o or "Сидоров" in o) and ("Ромашка" in o or "Лидер" in o or "Вектор" in o)) for o in orgs)
    w(f"  {'СКЛЕЕНО!!' if merged_two_parties else 'ok':9s} {text!r}")
    w(f"           ORG={orgs}")

w("\n========== C. Составные — ПОЗИТИВНЫЕ ==========")
POS = [
    "ИП Пирогова А.С. заключила договор",
    "с ИП Пироговой А.С. подписан акт",
    "Индивидуальный предприниматель Сидоров И.И.",
    "ИП Пирогова А.С.",
]
for text in POS:
    m = compound(text)
    orgs = [e.original_text for e in m if e.entity_type == "ORG"]
    joined = any(("Пирогов" in o or "Сидоров" in o) for o in orgs)
    w(f"  {'ok-склеено' if joined else 'НЕ СКЛЕЕНО':11s} {text!r} ORG={orgs}")

# ================= D. Batching equivalence (batched vs per-window ref) =================
w("\n========== D. Батчинг B3 — эквивалентность батч vs по-окно ==========")
def ref_boundary(doc, config_path):
    """Референс: по-оконная детекция (НЕ батч) той же геометрии окна."""
    import uuid as _uuid
    from tokenizer import _BOUNDARY_WINDOW, _boundary_sep
    segs = doc.segments
    extra = []
    for i in range(len(segs)-1):
        seg_a, seg_b = segs[i], segs[i+1]
        if not seg_a.text or not seg_b.text: continue
        ta, tb = seg_a.text, seg_b.text
        toff = max(0, len(ta)-_BOUNDARY_WINDOW)
        tail = ta[toff:]; head = tb[:_BOUNDARY_WINDOW]
        sep = _boundary_sep(seg_a, seg_b)
        win = tail+sep+head
        tail_end=len(tail); head_start=len(tail)+len(sep)
        wdoc = SourceDocument(segments=[TextSegment(id="w", text=win, source_type="txt_line", metadata={})], source_format="txt", source_path="x")
        wents = detect_regex(wdoc, config_path)+detect_ner(wdoc, config_path)
        for e in wents:
            ls, le = e.start, e.end
            if not (ls < tail_end and le > head_start): continue
            sa=toff+ls; ea=len(ta); eb=le-head_start
            extra.append((seg_a.id, sa, ea, e.entity_type, ta[sa:ea], e.detector))
            extra.append((seg_b.id, 0, eb, e.entity_type, tb[0:eb], e.detector))
    return sorted(extra)

# doc with boundary scenarios
segs = [
    TextSegment(id="a1", text="Телефон 8-900-", source_type="txt_line", metadata={}),
    TextSegment(id="a2", text="123-45-67 для связи", source_type="txt_line", metadata={}),
    TextSegment(id="b1", text="ИНН 7707", source_type="txt_line", metadata={}),
    TextSegment(id="b2", text="083893 организации", source_type="txt_line", metadata={}),
    TextSegment(id="c1", text="г. Москва, ул.", source_type="txt_line", metadata={}),
    TextSegment(id="c2", text="Ленина, д. 5", source_type="txt_line", metadata={}),
    TextSegment(id="d1", text="обычный сегмент", source_type="txt_line", metadata={}),
    TextSegment(id="d2", text="ещё один сегмент", source_type="txt_line", metadata={}),
]
bdoc = SourceDocument(segments=segs, source_format="txt", source_path="<b3>")
batched = sorted((e.segment_id, e.start, e.end, e.entity_type, e.original_text, e.detector)
                 for e in _detect_boundary_entities(bdoc, CFG))
ref = ref_boundary(bdoc, CFG)
w("батч   :", batched)
w("референс:", ref)
w("ИДЕНТИЧНЫ:", batched == ref)

with open(os.path.join(os.path.dirname(__file__), "bcde_out.txt"), "w", encoding="utf-8") as f:
    f.write(OUT.getvalue())
print("done")
