# -*- coding: utf-8 -*-
import os, sys, io
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "src"))
from docx import Document
from regex_detector import detect_regex
from ner_detector import detect_ner
from syntax_compound import merge_compound_entities
from tokenizer import tokenize
from extractor import extract
CFG = os.path.join(os.path.dirname(__file__), "..", "entity_types.yaml")
OUT=io.StringIO()
def w(*a): OUT.write(" ".join(str(x) for x in a)+"\n")
def enc(doc):
    ents=detect_regex(doc,CFG)+detect_ner(doc,CFG); ents=merge_compound_entities(doc,ents)
    return tokenize(doc,ents,CFG)

w("=== F2. Граница внутри ОДНОЙ СТРОКИ ТАБЛИЦЫ (sep='') + all_caps ===")
# split INN across two same-row all_caps cells
d=Document(); t=d.add_table(rows=1, cols=2)
for cell,txt in zip(t.rows[0].cells, ["инн 7707","083893"]):
    r=cell.paragraphs[0].add_run(txt); r.font.all_caps=True
p=os.path.join(os.path.dirname(__file__),"f2_inn.docx"); d.save(p)
doc=extract(p)
w("сегменты:", [(s.source_type, s.text, s.metadata.get('detection_text')) for s in doc.segments])
anon,final=enc(doc)
w("  anon:", repr(anon))
w("  утечка 7707083893 целиком:", "7707083893" in anon.replace(" ","").replace("\n","").replace("|",""))
w("  токены:", [(e.entity_type,e.original_text) for e in final])

with open(os.path.join(os.path.dirname(__file__),"f2_out.txt"),"w",encoding="utf-8") as f:
    f.write(OUT.getvalue())
print("done")
