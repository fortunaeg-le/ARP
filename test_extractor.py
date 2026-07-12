"""Тесты блока 1 — extractor.py (extract(path) -> SourceDocument)."""
import sys

import pytest
from docx import Document

from extractor import extract
from models import SourceDocument, TextSegment


class TestExtractDocxOrder:
    """Приёмка блока 1: сегменты покрывают тело документа в исходном порядке."""

    def test_docx_heading_paragraph_table_paragraph_segment_count(self, docx_factory):
        """Спека, блок 1, Приёмка: .docx с параграфами+таблицей 2x2+заголовком -> список сегментов покрывает весь текст тела."""
        path = docx_factory(
            heading="Договор поставки",
            paragraphs_before=["Преамбула договора."],
            table_rows=[["Строка1Кол1", "Строка1Кол2"], ["Строка2Кол1", "Строка2Кол2"]],
            paragraphs_after=["Заключение договора."],
        )
        doc = extract(path)
        assert len(doc.segments) == 7

    def test_docx_table_stays_between_surrounding_paragraphs_not_at_end(self, docx_factory):
        """Спека, блок 1, Приёмка: таблица стоит между окружающими параграфами, не в конце списка."""
        path = docx_factory(
            heading="Договор поставки",
            paragraphs_before=["Преамбула договора."],
            table_rows=[["A1", "B1"], ["A2", "B2"]],
            paragraphs_after=["Заключение договора."],
        )
        doc = extract(path)
        ids_in_order = [s.id for s in doc.segments]
        assert ids_in_order == ["p0", "p1", "t0_r0_c0", "t0_r0_c1", "t0_r1_c0", "t0_r1_c1", "p2"]

    def test_docx_segment_ids_are_unique(self, docx_factory):
        """HANDOFF_1, Инварианты: TextSegment.id уникален в пределах SourceDocument.segments."""
        path = docx_factory(
            heading="H",
            paragraphs_before=["A"],
            table_rows=[["1", "2"], ["3", "4"]],
            paragraphs_after=["B"],
        )
        doc = extract(path)
        ids = [s.id for s in doc.segments]
        assert len(ids) == len(set(ids))

    def test_docx_heading_style_matches_word_style_name(self, docx_factory):
        """Спека, блок 1: metadata['style'] для заголовка непусто и совпадает с именем стиля Word."""
        path = docx_factory(heading="Договор поставки", paragraphs_before=[], paragraphs_after=[])
        doc = extract(path)
        heading_seg = doc.segments[0]
        assert heading_seg.metadata["style"] == "Heading 1"

    def test_docx_normal_paragraph_style_is_normal(self, docx_factory):
        """Спека, блок 1: metadata['style'] записывается как есть (paragraph.style.name), без нормализации."""
        path = docx_factory(paragraphs_before=["Преамбула договора."])
        doc = extract(path)
        assert doc.segments[0].metadata["style"] == "Normal"

    def test_docx_paragraph_source_type_is_docx_paragraph(self, docx_factory):
        """Спека, блок 1: параграфы получают source_type='docx_paragraph' явно."""
        path = docx_factory(paragraphs_before=["Текст."])
        doc = extract(path)
        assert doc.segments[0].source_type == "docx_paragraph"

    def test_docx_table_cell_source_type_is_docx_table_cell(self, docx_factory):
        """Спека, блок 1: ячейки таблицы получают source_type='docx_table_cell'."""
        path = docx_factory(table_rows=[["A", "B"]])
        doc = extract(path)
        assert doc.segments[0].source_type == "docx_table_cell"

    def test_docx_table_cell_metadata_has_all_three_keys(self, docx_factory):
        """Спека, блок 1: metadata ячейки обязательно содержит все три ключа table_index/row_index/col_index."""
        path = docx_factory(table_rows=[["A", "B"], ["C", "D"]])
        doc = extract(path)
        cell = doc.segments[0]
        assert set(cell.metadata.keys()) == {"table_index", "row_index", "col_index"}

    def test_docx_paragraph_metadata_has_index_and_style_keys(self, docx_factory):
        """HANDOFF_1, Инварианты: metadata параграфа содержит ключи paragraph_index и style."""
        path = docx_factory(paragraphs_before=["Текст."])
        doc = extract(path)
        assert set(doc.segments[0].metadata.keys()) == {"paragraph_index", "style"}

    def test_docx_source_format_is_docx(self, docx_factory):
        """Спека, блок 1: SourceDocument.source_format == 'docx' для .docx входа."""
        path = docx_factory(paragraphs_before=["Текст."])
        doc = extract(path)
        assert doc.source_format == "docx"

    def test_docx_source_path_matches_input(self, docx_factory):
        """Спека, блок 1: SourceDocument.source_path равен переданному пути."""
        path = docx_factory(paragraphs_before=["Текст."])
        doc = extract(path)
        assert doc.source_path == path


class TestExtractDocxEmptySegments:
    """Спека, блок 1: пустые параграфы/строки включаются в список, не пропускаются."""

    def test_docx_empty_paragraph_is_included_with_empty_text(self, docx_factory):
        """Спека, блок 1: пустые параграфы — включать в список сегментов (с пустым text), не пропускать."""
        path = docx_factory(paragraphs_before=["Текст.", "", "Ещё текст."])
        doc = extract(path)
        assert len(doc.segments) == 3
        assert doc.segments[1].text == ""


class TestExtractDocxMergedCells:
    """HANDOFF_1, Пример 2 — merged-ячейка: координата сохраняется, текст не дублируется."""

    def test_merged_cell_keeps_text_in_first_coordinate_only(self, tmp_path):
        """HANDOFF_1, Пример 2: t0_r0_c0 содержит 'Объединённая ячейка', t0_r0_c1 — пустая строка."""
        path = tmp_path / "merged.docx"
        document = Document()
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).merge(table.cell(0, 1))
        table.cell(0, 0).text = "Объединённая ячейка"
        table.cell(1, 0).text = "Строка2Кол1"
        table.cell(1, 1).text = "Строка2Кол2"
        document.save(path)

        doc = extract(str(path))
        by_id = {s.id: s for s in doc.segments}
        assert by_id["t0_r0_c0"].text == "Объединённая ячейка"
        assert by_id["t0_r0_c1"].text == ""

    def test_merged_cell_coordinate_grid_is_preserved_not_dropped(self, tmp_path):
        """Спека, блок 1: merge не должен схлопывать координатную сетку — все 4 сегмента присутствуют."""
        path = tmp_path / "merged.docx"
        document = Document()
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).merge(table.cell(0, 1))
        document.save(path)

        doc = extract(str(path))
        ids = {s.id for s in doc.segments}
        assert ids == {"t0_r0_c0", "t0_r0_c1", "t0_r1_c0", "t0_r1_c1"}


class TestExtractDocxNestedTableWarning:
    """Спека, блок 1: вложенная таблица в ячейке — предупреждение в stderr, не извлекается."""

    def test_nested_table_prints_warning_with_cell_coordinates(self, tmp_path, capsys):
        """Спека, блок 1: обнаружена вложенная таблица -> предупреждение в stderr с координатами ячейки."""
        path = tmp_path / "nested.docx"
        document = Document()
        outer = document.add_table(rows=1, cols=1)
        outer.cell(0, 0).add_table(rows=1, cols=1)
        document.save(path)

        extract(str(path))
        captured = capsys.readouterr()
        assert "t0_r0_c0" in captured.err


class TestExtractTxtBasic:
    """HANDOFF_1, Пример 3 — .txt граничный случай."""

    def test_txt_crlf_lines_and_trailing_newline_boundary(self, txt_factory):
        """HANDOFF_1, Пример 3: файл '...\\r\\n\\r\\nЧетвёртая\\r\\n' -> 4 сегмента, без 5-го пустого от хвостового \\r\\n."""
        content = "Первая строка\r\nВторая строка\r\n\r\nЧетвёртая\r\n"
        path = txt_factory(content, encoding="utf-8-sig")
        doc = extract(path)
        texts = [s.text for s in doc.segments]
        assert texts == ["Первая строка", "Вторая строка", "", "Четвёртая"]

    def test_txt_line_ids_are_l_prefixed_sequential(self, txt_factory):
        """Спека, блок 1: id строки txt вида l{index}."""
        content = "Первая\r\nВторая\r\n"
        path = txt_factory(content)
        doc = extract(path)
        assert [s.id for s in doc.segments] == ["l0", "l1"]

    def test_txt_source_type_is_txt_line(self, txt_factory):
        """Спека, блок 1: source_type == 'txt_line' для строк .txt."""
        path = txt_factory("Строка\r\n")
        doc = extract(path)
        assert doc.segments[0].source_type == "txt_line"

    def test_txt_metadata_has_line_index_and_encoding_keys(self, txt_factory):
        """Спека, блок 1: metadata строки .txt содержит оба ключа line_index и encoding."""
        path = txt_factory("Строка\r\n")
        doc = extract(path)
        assert set(doc.segments[0].metadata.keys()) == {"line_index", "encoding"}

    def test_txt_encoding_metadata_is_utf8_sig_when_that_encoding_succeeds(self, txt_factory):
        """HANDOFF_1, Пример 3: encoding в metadata == 'utf-8-sig', когда файл в этой кодировке."""
        path = txt_factory("Текст\r\n", encoding="utf-8-sig")
        doc = extract(path)
        assert doc.segments[0].metadata["encoding"] == "utf-8-sig"

    def test_txt_source_format_is_txt(self, txt_factory):
        """Спека, блок 1: SourceDocument.source_format == 'txt' для .txt входа."""
        path = txt_factory("Текст\r\n")
        doc = extract(path)
        assert doc.source_format == "txt"

    def test_txt_no_trailing_newline_does_not_add_extra_empty_segment(self, txt_factory):
        """Спека, блок 1: если файл НЕ заканчивается на \\n, финальный пустой элемент не отбрасывается лишний раз."""
        path = txt_factory("Первая\r\nВторая", encoding="utf-8-sig")
        doc = extract(path)
        assert [s.text for s in doc.segments] == ["Первая", "Вторая"]

    def test_txt_minimal_single_empty_line_input(self, txt_factory):
        """Граничный случай: пустой .txt файл (0 байт) даёт 0 сегментов (пустая строка отбрасывается как хвостовая)."""
        path = txt_factory("")
        doc = extract(path)
        assert doc.segments == []


class TestExtractTxtEncodingFallback:
    """Спека, блок 1: при UnicodeDecodeError на utf-8-sig — повторная попытка с cp1251."""

    def test_cp1251_file_falls_back_and_decodes_correctly(self, tmp_path):
        """Спека, блок 1: .txt в cp1251 (невалиден как utf-8-sig) -> корректно декодируется через cp1251 fallback."""
        path = tmp_path / "cp1251.txt"
        with open(path, "w", encoding="cp1251", newline="") as f:
            f.write("Русский текст в cp1251\r\n")
        doc = extract(str(path))
        assert doc.segments[0].text == "Русский текст в cp1251"
        assert doc.segments[0].metadata["encoding"] == "cp1251"


class TestExtractErrors:
    """HANDOFF_1, Пример 4 — ошибки extract()."""

    def test_missing_file_raises_file_not_found_error(self):
        """HANDOFF_1, Пример 4: extract('nope.docx') -> FileNotFoundError."""
        with pytest.raises(FileNotFoundError):
            extract("nope_does_not_exist.docx")

    def test_unsupported_extension_raises_value_error_with_message(self, tmp_path):
        """HANDOFF_1, Пример 4: extract('file.pdf') -> ValueError('Неподдерживаемый формат: .pdf. Поддерживаются: .docx, .txt')."""
        path = tmp_path / "file.pdf"
        path.write_bytes(b"dummy")
        with pytest.raises(ValueError, match=r"Неподдерживаемый формат: \.pdf\. Поддерживаются: \.docx, \.txt"):
            extract(str(path))

    def test_unsupported_extension_is_case_insensitive(self, tmp_path):
        """Спека, блок 1: расширение определяется в нижнем регистре."""
        path = tmp_path / "file.PDF"
        path.write_bytes(b"dummy")
        with pytest.raises(ValueError):
            extract(str(path))
